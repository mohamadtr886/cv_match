import streamlit as st
import pandas as pd
from io import BytesIO
import docx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
import PyPDF2
import re
import json
from collections import Counter

# --- PAGE CONFIGURATION ---
st.set_page_config(
    page_title="CV Match AI",
    page_icon="💼",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# Hide Streamlit chrome — watermark, toolbar, deploy button
st.markdown(
    "<style>"
    "#MainMenu,footer,header{visibility:hidden!important;}"
    "[data-testid='stToolbar'],[data-testid='stDecoration'],"
    "[data-testid='stStatusWidget'],.stDeployButton{display:none!important;}"
    "</style>",
    unsafe_allow_html=True,
)

# --- INITIALIZE SESSION STATE ---
defaults = {
    'active_tab': 0,
    'job_role': "",
    'cv_full_text': "",
    'job_desc': "",
    'input_method': "Upload File",
    'manual_cv_data': {},
    'follow_up_answers': {},        # {key: {question, answer, section, gap_addressed}}
    'analysis_results': {"matches": [], "missing": [], "score": 0},
    'ai_cv_analysis': {},           # Step 1+2: deep CV+JD analysis from Claude
    'dynamic_questions': [],        # Step 3: personalized questions [{key, question, hint, section, gap_addressed}]
    'improvement_log': [],          # Step 10: [{type, section, description}]
    'final_cv_data': {},
    'adjusted_cv_data': {},
    'active_lang': "English",
    'cv_uploaded': False,
    'analysis_done': False,
    'cv_adjusted': False,
    'show_landing': True,
}
for k, v in defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v

# --- CUSTOM CSS ---
st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:ital,wght@0,400;0,500;0,600;0,700;0,800;1,400&display=swap');

    html, body, [class*="css"] {
        font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif !important;
    }

    /* ── DESIGN TOKENS ── */
    :root {
        --blue:        #2563EB;
        --blue-hover:  #1D4ED8;
        --blue-light:  #EFF6FF;
        --blue-mid:    #BFDBFE;
        --text:        #0F172A;
        --muted:       #64748B;
        --border:      #E5E7EB;
        --bg:          #F8FAFC;
        --card:        #FFFFFF;
        --success:     #16A34A;
        --warning:     #F59E0B;
        --error:       #DC2626;
    }

    /* ── GLOBAL RESETS ── */
    .block-container {
        padding-top: 0 !important;
        padding-left: 2rem !important;
        padding-right: 2rem !important;
        max-width: 880px !important;
        background: var(--bg) !important;
    }
    .stApp { background: var(--bg) !important; }
    #MainMenu, footer, header { visibility: hidden; }
    div[data-testid="stTabs"] { display: none; }

    /* ── STEP BAR ── */
    .rp-step-bar {
        display: flex;
        align-items: center;
        background: var(--card);
        border: 1px solid var(--border);
        border-radius: 12px;
        padding: 12px 18px;
        margin-bottom: 28px;
        box-shadow: 0 1px 3px rgba(0,0,0,0.04);
        gap: 0;
    }
    .rp-step-item { display: flex; align-items: center; gap: 8px; flex: 1; }
    .rp-step-dot {
        width: 22px; height: 22px; border-radius: 50%;
        display: flex; align-items: center; justify-content: center;
        font-size: 10px; font-weight: 700; flex-shrink: 0;
        transition: all 0.2s;
    }
    .rp-dot-done   { background: var(--blue); color: white; }
    .rp-dot-active { background: var(--text); color: white; }
    .rp-dot-todo   { background: var(--border); color: #9CA3AF; }
    .rp-step-name  { font-size: 12px; font-weight: 500; }
    .rp-name-done  { color: var(--blue); font-weight: 600; }
    .rp-name-active{ color: var(--text); font-weight: 700; }
    .rp-name-todo  { color: #9CA3AF; }
    .rp-connector  { height: 2px; width: 20px; flex-shrink: 0; }
    .rp-conn-done  { background: var(--blue); }
    .rp-conn-todo  { background: var(--border); }

    /* ── LOGO / HEADER ── */
    .rp-header {
        display: flex;
        align-items: center;
        gap: 10px;
        padding: 18px 0 16px;
        border-bottom: 1px solid var(--border);
        margin-bottom: 28px;
    }
    .rp-brand   { font-size: 17px; font-weight: 700; letter-spacing: -0.4px; color: var(--text); }
    .rp-tagline { font-size: 13px; color: var(--muted); margin-left: auto; }

    /* ── PAGE TITLES ── */
    .rp-page-badge { font-size: 11px; color: var(--blue); font-weight: 600; letter-spacing: 0.6px; margin-bottom: 6px; text-transform: uppercase; }
    .rp-page-title { font-size: 26px; font-weight: 700; letter-spacing: -0.6px; color: var(--text); margin-bottom: 5px; }
    .rp-page-sub   { font-size: 13px; color: var(--muted); margin-bottom: 24px; line-height: 1.6; }

    /* ── CARDS ── */
    .rp-card, .rp-card-white, .content-card {
        background: var(--card);
        border: 1px solid var(--border);
        border-radius: 12px;
        padding: 16px 20px;
        margin-bottom: 16px;
        box-shadow: 0 1px 4px rgba(0,0,0,0.04);
    }
    .rp-card { background: var(--bg); }

    /* ── PILLS ── */
    .rp-pill { font-size: 11px; font-weight: 600; padding: 2px 10px; border-radius: 9999px; }
    .rp-pill-strong  { background: #DCFCE7; color: #166534; }
    .rp-pill-partial { background: #FEF9C3; color: #92400E; }
    .rp-pill-missing { background: #FEE2E2; color: #991B1B; }

    /* ── TIP BOX ── */
    .rp-tip {
        background: var(--blue-light);
        border: 1px solid var(--blue-mid);
        border-radius: 10px;
        padding: 14px 16px;
        font-size: 13px;
        color: #1E40AF;
        line-height: 1.6;
        margin-top: 14px;
    }

    /* ── REFINEMENT QUESTION NUM ── */
    .rp-q-num {
        display: inline-flex; align-items: center; justify-content: center;
        background: var(--blue); color: white;
        width: 20px; height: 20px; border-radius: 50%;
        font-size: 10px; font-weight: 700; margin-right: 6px;
    }
    .rp-q-text { font-size: 13px; font-weight: 500; color: var(--text); margin-bottom: 8px; }

    /* ── DOWNLOAD CARDS ── */
    .rp-dl-fmt  { font-weight: 700; font-size: 15px; margin-bottom: 3px; color: var(--text); }
    .rp-dl-desc { font-size: 12px; color: var(--muted); margin-bottom: 14px; }

    /* ── CHECKLIST ── */
    .rp-checklist { border: 1px solid var(--border); border-radius: 12px; padding: 18px 20px; background: var(--card); box-shadow: 0 1px 3px rgba(0,0,0,0.04); }
    .rp-checklist-title { font-weight: 600; font-size: 13px; margin-bottom: 10px; color: var(--text); }
    .rp-check-item { font-size: 12px; color: var(--muted); padding: 5px 0; border-bottom: 1px solid var(--border); }
    .rp-check-item:last-child { border-bottom: none; }

    /* ── ALERT BOXES ── */
    .alert-box { padding: 12px 16px; border-radius: 8px; font-size: 13px; font-weight: 500; margin-bottom: 12px; }
    .alert-green  { background: #DCFCE7; color: #166534; }
    .alert-orange { background: #FEF3C7; color: #92400E; }
    .alert-red    { background: #FEE2E2; color: #991B1B; }

    /* ── PRIMARY BUTTON ── */
    .stButton > button, .stDownloadButton > button {
        border-radius: 8px !important;
        font-weight: 600 !important;
        font-size: 13px !important;
        font-family: 'Inter', sans-serif !important;
        transition: all 0.15s !important;
        padding: 8px 18px !important;
    }
    .stButton > button[kind="primary"],
    .stDownloadButton > button[kind="primary"] {
        background: var(--blue) !important;
        color: white !important;
        border: none !important;
        box-shadow: 0 1px 3px rgba(37,99,235,0.3) !important;
    }
    .stButton > button[kind="primary"]:hover,
    .stDownloadButton > button[kind="primary"]:hover {
        background: var(--blue-hover) !important;
        box-shadow: 0 2px 8px rgba(37,99,235,0.4) !important;
    }
    .stButton > button[kind="secondary"],
    .stDownloadButton > button[kind="secondary"] {
        background: transparent !important;
        color: var(--blue) !important;
        border: 1.5px solid var(--blue) !important;
    }
    .stButton > button[kind="secondary"]:hover,
    .stDownloadButton > button[kind="secondary"]:hover {
        background: var(--blue-light) !important;
    }
    /* Disabled button */
    .stButton > button:disabled {
        background: #E5E7EB !important;
        color: #9CA3AF !important;
        border: none !important;
        box-shadow: none !important;
        cursor: not-allowed !important;
    }

    /* ── INPUT FIELDS ── */
    .stTextInput input, .stTextArea textarea {
        border-radius: 8px !important;
        border: 1.5px solid var(--border) !important;
        background: var(--card) !important;
        font-size: 13px !important;
        font-family: 'Inter', sans-serif !important;
        color: var(--text) !important;
        transition: border-color 0.15s !important;
    }
    .stTextInput input:focus, .stTextArea textarea:focus {
        border-color: var(--blue) !important;
        box-shadow: 0 0 0 3px rgba(37,99,235,0.1) !important;
        background: white !important;
    }
    .stTextInput label, .stTextArea label, .stFileUploader label {
        font-size: 13px !important;
        font-weight: 600 !important;
        color: var(--text) !important;
    }

    /* ── RADIO AS SELECTABLE CARDS ── */
    div[data-testid="stRadio"] > label {
        font-size: 13px !important;
        font-weight: 600 !important;
        color: var(--text) !important;
        margin-bottom: 10px !important;
    }
    div[data-testid="stRadio"] > div[role="radiogroup"] {
        display: flex !important;
        flex-direction: row !important;
        gap: 10px !important;
        flex-wrap: wrap !important;
    }
    div[data-testid="stRadio"] > div[role="radiogroup"] > label {
        display: flex !important;
        align-items: center !important;
        gap: 8px !important;
        background: var(--card) !important;
        border: 1.5px solid var(--border) !important;
        border-radius: 8px !important;
        padding: 10px 16px !important;
        cursor: pointer !important;
        transition: all 0.15s !important;
        font-size: 13px !important;
        font-weight: 500 !important;
        color: var(--text) !important;
        min-width: 130px !important;
    }
    div[data-testid="stRadio"] > div[role="radiogroup"] > label:has(input:checked) {
        border-color: var(--blue) !important;
        background: var(--blue-light) !important;
        color: var(--blue) !important;
        font-weight: 600 !important;
    }
    div[data-testid="stRadio"] > div[role="radiogroup"] > label:hover {
        border-color: var(--blue) !important;
    }

    /* ── EXPANDER ── */
    .streamlit-expanderHeader {
        font-size: 13px !important;
        font-weight: 600 !important;
        color: var(--text) !important;
        border-radius: 8px !important;
        background: var(--card) !important;
        border: 1px solid var(--border) !important;
    }

    /* ── STREAMLIT INFO / WARNING / ERROR ── */
    div[data-testid="stNotification"] {
        border-radius: 10px !important;
    }

    /* ── SCORE DARK CARD ── */
    .rp-score-dark {
        background: var(--text);
        color: white;
        border-radius: 12px;
        padding: 22px 26px;
        display: flex; align-items: center; justify-content: space-between;
        margin-bottom: 14px;
    }
    .rp-score-num   { font-size: 48px; font-weight: 800; letter-spacing: -2px; line-height: 1; }
    .rp-score-label { font-size: 11px; color: rgba(255,255,255,0.5); margin-bottom: 3px; letter-spacing: 0.4px; }
    .rp-score-desc  { font-size: 12px; color: rgba(255,255,255,0.45); margin-top: 4px; }
    .rp-score-right { text-align: right; font-size: 12px; line-height: 2.2; }
    .rp-score-s { color: #86EFAC; }
    .rp-score-p { color: rgba(255,255,255,0.5); }
    .rp-score-m { color: rgba(255,255,255,0.3); }

    /* ── ANALYSIS ITEMS ── */
    .rp-analysis-item {
        display: flex; align-items: center; justify-content: space-between;
        padding: 9px 13px; border-radius: 8px;
        background: var(--bg); border: 1px solid var(--border);
        margin-bottom: 5px; font-size: 13px; color: var(--text);
    }

    /* ── LANDING PAGE ── */
    .rp-hero { max-width: 640px; margin: 0 auto; padding: 80px 32px 48px; text-align: center; }
    .rp-hero h1 {
        font-size: 50px; font-weight: 800; line-height: 1.1;
        letter-spacing: -2px; margin-bottom: 20px; color: var(--text);
    }
    .rp-hero h1 span { color: var(--blue); }
    .rp-hero p {
        font-size: 16px; color: var(--muted); line-height: 1.7;
        max-width: 460px; margin: 0 auto 36px;
    }

    /* ── BENTO GRID ── */
    .rp-bento {
        max-width: 840px; margin: 40px auto 0; padding: 0 32px 56px;
        display: grid;
        grid-template-columns: repeat(3, 1fr);
        grid-template-rows: auto auto;
        gap: 12px;
    }
    .rp-bento-cell {
        background: var(--card);
        border: 1px solid var(--border);
        border-radius: 14px;
        padding: 22px;
        position: relative; overflow: hidden;
        transition: box-shadow 0.2s, transform 0.2s;
    }
    .rp-bento-cell:hover {
        box-shadow: 0 6px 24px rgba(37,99,235,0.08);
        transform: translateY(-1px);
    }
    .rp-bento-cell::after {
        content: '';
        position: absolute; top: 0; left: 0; right: 0; height: 3px;
        background: var(--blue);
        transform: scaleX(0); transform-origin: left;
        transition: transform 0.25s;
    }
    .rp-bento-cell:hover::after { transform: scaleX(1); }
    .rp-bento-icon  { font-size: 20px; margin-bottom: 10px; }
    .rp-bento-title { font-size: 13px; font-weight: 700; color: var(--text); margin-bottom: 5px; }
    .rp-bento-desc  { font-size: 12px; color: var(--muted); line-height: 1.55; }

    /* ── HOW IT WORKS ── */
    .rp-how {
        background: var(--bg);
        border-top: 1px solid var(--border); border-bottom: 1px solid var(--border);
        padding: 60px 32px;
    }
    .rp-how-inner  { max-width: 700px; margin: 0 auto; }
    .rp-how-label  { font-size: 11px; font-weight: 600; color: var(--blue); letter-spacing: 1px; text-align: center; margin-bottom: 10px; text-transform: uppercase; }
    .rp-how-title  { font-size: 28px; font-weight: 700; letter-spacing: -0.8px; text-align: center; margin-bottom: 40px; color: var(--text); }
    .rp-how-steps  { display: grid; grid-template-columns: repeat(3,1fr); gap: 28px; }
    .rp-how-step   { text-align: center; }
    .rp-how-num    { width: 34px; height: 34px; border-radius: 50%; background: var(--blue); color: white; font-size: 13px; font-weight: 700; display: flex; align-items: center; justify-content: center; margin: 0 auto 12px; }
    .rp-how-step-title { font-size: 14px; font-weight: 600; margin-bottom: 5px; color: var(--text); }
    .rp-how-step-desc  { font-size: 12px; color: var(--muted); line-height: 1.55; }

    /* ── CTA + FOOTER ── */
    .rp-cta { max-width: 560px; margin: 0 auto; padding: 64px 32px; text-align: center; }
    .rp-cta h2 { font-size: 30px; font-weight: 700; letter-spacing: -0.8px; margin-bottom: 12px; color: var(--text); }
    .rp-cta p  { font-size: 14px; color: var(--muted); margin-bottom: 24px; }
    .rp-footer { border-top: 1px solid var(--border); padding: 20px 48px; display: flex; align-items: center; justify-content: space-between; }
    .rp-footer-brand { font-weight: 700; font-size: 14px; color: var(--text); }
    .rp-footer-p     { font-size: 12px; color: var(--muted); }

    /* ── SCORE ROW ON LANDING ── */
    .rp-score-row {
        display: flex; align-items: center; justify-content: space-between;
        padding: 7px 11px; border-radius: 6px;
        background: var(--bg); border: 1px solid var(--border);
        margin-bottom: 4px; font-size: 13px;
    }
    </style>
""", unsafe_allow_html=True)

# ============================================================
# CONSTANTS & DICTIONARIES (unchanged from original)
# ============================================================

STOPWORDS = {
    "a","an","the","and","or","but","if","then","else","when","at","from","by",
    "for","with","about","against","between","into","through","during","before",
    "after","above","below","to","up","down","in","out","on","off","over","under",
    "again","further","once","here","there","all","any","both","each","few","more",
    "most","other","some","such","no","nor","not","only","own","same","so","than",
    "too","very","s","t","can","will","just","don","should","now","i","me","my",
    "myself","we","our","ours","ourselves","you","your","yours","yourself",
    "yourselves","he","him","his","himself","she","her","hers","herself","it","its",
    "itself","they","them","their","theirs","themselves","what","which","who","whom",
    "this","that","these","those","am","is","are","was","were","be","been","being",
    "have","has","had","having","do","does","did","doing"
}

SKILL_KEYWORDS = {
    "python","sql","java","javascript","react","node","html","css","git","github",
    "r","matlab","tableau","powerbi","power bi","excel","bigquery","pandas","numpy",
    "scikit-learn","tensorflow","pytorch","aws","azure","docker","kubernetes","unix",
    "linux","bash","spark","hadoop","c++","c#","php","ruby","swift","kotlin",
    "typescript","scala","go","rust","mongodb","postgresql","mysql","redis","kafka",
    "airflow","dbt","looker","databricks","snowflake","jira","confluence","jenkins",
}

# ============================================================
# SKILL GROUPS — Heart of the hybrid matching engine
# Each group has: trigger patterns (what to look for in JD),
# strong keywords (direct CV match), partial keywords (related),
# weight (importance), and a display label.
# ============================================================

SKILL_GROUPS = [
    # ── Programming Languages ──
    {
        "label": "Python",
        "weight": 3,
        "jd_patterns": [r'\bpython\b'],
        "strong": ["python", "pandas", "numpy", "scipy", "flask", "django", "scikit"],
        "partial": ["programming", "scripting", "automation", "jupyter"],
        "category": "Technical"
    },
    {
        "label": "SQL / Databases",
        "weight": 3,
        "jd_patterns": [r'\bsql\b', r'\bdatabase', r'\bqueries\b'],
        "strong": ["sql", "mysql", "postgresql", "sqlite", "bigquery", "database", "queries"],
        "partial": ["data manipulation", "data extraction", "spreadsheet", "excel"],
        "category": "Technical"
    },
    {
        "label": "R / Statistical Computing",
        "weight": 2,
        "jd_patterns": [r'\br\b(?:\s+programming|\s+language|\s+studio)?', r'\brstudio\b', r'\bggplot\b'],
        "strong": ["r", "rstudio", "ggplot", "tidyverse", "dplyr", "r programming"],
        "partial": ["statistics", "statistical analysis", "spss", "matlab", "python"],
        "category": "Technical"
    },
    {
        "label": "Data Visualization",
        "weight": 2,
        "jd_patterns": [r'\bvisuali', r'\btableau\b', r'\bpower\s*bi\b', r'\bdashboard'],
        "strong": ["tableau", "power bi", "powerbi", "matplotlib", "plotly", "seaborn", "ggplot", "dashboard", "visualization"],
        "partial": ["charts", "graphs", "reports", "excel", "presentations", "plots"],
        "category": "Technical"
    },
    {
        "label": "Machine Learning",
        "weight": 2,
        "jd_patterns": [r'\bmachine\s+learning\b', r'\bml\b', r'\bai\b', r'\bclassif', r'\bcluster'],
        "strong": ["machine learning", "ml", "scikit-learn", "sklearn", "tensorflow", "pytorch", "classification", "clustering", "neural"],
        "partial": ["predictive", "modeling", "statistical modeling", "regression", "deep learning", "ai"],
        "category": "Technical"
    },
    {
        "label": "Statistical Analysis",
        "weight": 2,
        "jd_patterns": [r'\bstatistic', r'\bregression\b', r'\bhypothesis', r'\bmodeling\b|\bmodelling\b'],
        "strong": ["statistics", "statistical", "regression", "hypothesis testing", "modeling", "quantitative", "spss", "anova"],
        "partial": ["data analysis", "analytics", "r", "python", "research", "math"],
        "category": "Technical"
    },
    {
        "label": "Git / Version Control",
        "weight": 1,
        "jd_patterns": [r'\bgit\b', r'\bversion\s+control\b', r'\bgithub\b'],
        "strong": ["git", "github", "gitlab", "version control", "bitbucket"],
        "partial": ["code", "repository", "collaboration", "open source"],
        "category": "Technical"
    },
    {
        "label": "Linux / Unix",
        "weight": 1,
        "jd_patterns": [r'\blinux\b', r'\bunix\b', r'\bbash\b', r'\bshell\b'],
        "strong": ["linux", "unix", "bash", "shell", "terminal", "command line"],
        "partial": ["scripting", "automation", "python", "programming"],
        "category": "Technical"
    },
    {
        "label": "Cloud (AWS / Azure / GCP)",
        "weight": 2,
        "jd_patterns": [r'\baws\b', r'\bazure\b', r'\bgcp\b', r'\bcloud\b'],
        "strong": ["aws", "azure", "gcp", "cloud", "s3", "lambda", "google cloud"],
        "partial": ["deployment", "docker", "kubernetes", "devops"],
        "category": "Technical"
    },
    {
        "label": "Excel / Spreadsheets",
        "weight": 1,
        "jd_patterns": [r'\bexcel\b', r'\bspreadsheet\b'],
        "strong": ["excel", "spreadsheet", "pivot", "vlookup", "google sheets"],
        "partial": ["data analysis", "reporting", "sql", "tables"],
        "category": "Technical"
    },
    {
        "label": "NLP / Text Analysis",
        "weight": 2,
        "jd_patterns": [r'\bnlp\b', r'\bnatural\s+language', r'\btext\s+(?:mining|analysis|processing)'],
        "strong": ["nlp", "natural language processing", "spacy", "nltk", "bert", "transformers", "text mining"],
        "partial": ["python", "machine learning", "classification", "text"],
        "category": "Technical"
    },
    {
        "label": "Network / Graph Analysis",
        "weight": 2,
        "jd_patterns": [r'\bnetwork\s+analysis\b', r'\bgraph\s+(?:theory|analysis)\b', r'\bnetworkx\b'],
        "strong": ["network analysis", "networkx", "graph", "nodes", "edges", "clustering coefficient"],
        "partial": ["python", "data analysis", "algorithms"],
        "category": "Technical"
    },
    {
        "label": "Agile / Scrum",
        "weight": 1,
        "jd_patterns": [r'\bagile\b', r'\bscrum\b', r'\bsprint\b', r'\bkanban\b'],
        "strong": ["agile", "scrum", "sprint", "kanban", "jira"],
        "partial": ["project management", "teamwork", "collaboration"],
        "category": "Technical"
    },
    # ── Education ──
    {
        "label": "Bachelor's Degree",
        "weight": 2,
        "jd_patterns": [r'\bbachelor', r'\bb\.?sc\b', r'\bundergraduate\b', r'\bdegree\b'],
        "strong": ["bachelor", "bsc", "b.sc", "undergraduate", "university", "college", "degree"],
        "partial": ["studying", "student", "final year", "graduate"],
        "category": "Education"
    },
    {
        "label": "Master's Degree",
        "weight": 2,
        "jd_patterns": [r'\bmaster', r'\bm\.?sc\b', r'\bmba\b', r'\bpostgraduate\b'],
        "strong": ["master", "msc", "m.sc", "mba", "postgraduate"],
        "partial": ["bachelor", "university", "advanced degree"],
        "category": "Education"
    },
    {
        "label": "Computer Science Degree",
        "weight": 2,
        "jd_patterns": [r'\bcomputer\s+science\b', r'\bcs\s+degree\b', r'\bsoftware\s+engineering\s+degree\b'],
        "strong": ["computer science", "cs degree", "software engineering", "information systems"],
        "partial": ["programming", "algorithms", "data structures", "university"],
        "category": "Education"
    },
    {
        "label": "Statistics / Math Degree",
        "weight": 2,
        "jd_patterns": [r'\bstatistics\s+degree\b', r'\bmathematic', r'\bquantitative\s+field\b'],
        "strong": ["statistics", "mathematics", "math", "quantitative", "data science degree"],
        "partial": ["engineering", "physics", "economics", "university"],
        "category": "Education"
    },
    # ── Languages ──
    {
        "label": "English",
        "weight": 1,
        "jd_patterns": [r'\benglish\b'],
        "strong": ["english"],
        "partial": ["fluent", "bilingual", "native"],
        "category": "Languages"
    },
    {
        "label": "Hebrew",
        "weight": 2,
        "jd_patterns": [r'\bhebrew\b'],
        "strong": ["hebrew"],
        "partial": ["fluent", "native", "israel"],
        "category": "Languages"
    },
    {
        "label": "Arabic",
        "weight": 1,
        "jd_patterns": [r'\barabic\b'],
        "strong": ["arabic"],
        "partial": ["fluent", "native"],
        "category": "Languages"
    },
    # ── Soft Skills ──
    {
        "label": "Communication Skills",
        "weight": 1,
        "jd_patterns": [r'\bcommunication\b', r'\bpresent', r'\bstakeholder'],
        "strong": ["communication", "presented", "presentation", "stakeholders", "written", "verbal", "public speaking"],
        "partial": ["tutoring", "teaching", "teamwork", "reporting", "documentation", "mentoring"],
        "category": "Soft Skills"
    },
    {
        "label": "Teamwork / Collaboration",
        "weight": 1,
        "jd_patterns": [r'\bteamwork\b', r'\bcollaborat', r'\bteam\s+player\b', r'\bcross.functional\b'],
        "strong": ["teamwork", "collaboration", "team player", "cross-functional", "worked with team"],
        "partial": ["project", "group", "together", "hackathon", "volunteer", "committee"],
        "category": "Soft Skills"
    },
    {
        "label": "Problem Solving",
        "weight": 1,
        "jd_patterns": [r'\bproblem.solv', r'\banalytical\b', r'\bcritical\s+think'],
        "strong": ["problem solving", "analytical", "critical thinking", "solutions"],
        "partial": ["research", "debugging", "algorithms", "optimized", "improved"],
        "category": "Soft Skills"
    },
    {
        "label": "Project Management",
        "weight": 1,
        "jd_patterns": [r'\bproject\s+management\b', r'\bproject\s+lead', r'\bcoordinat'],
        "strong": ["project management", "project lead", "coordinated", "managed project"],
        "partial": ["organized", "planned", "delivered", "timeline", "agile"],
        "category": "Soft Skills"
    },
]

# Words to NEVER treat as requirements — noise filter
NOISE_WORDS = {
    "junior","senior","analyst","engineer","developer","manager","associate","intern",
    "looking","seeking","hiring","join","growing","excited","dynamic","innovative",
    "responsibilities","requirements","qualifications","preferred","required","needed",
    "must","should","will","would","please","submit","apply","resume","cv",
    "company","team","organization","department","role","position","opportunity",
    "years","experience","background","foundation","ability","understanding",
    "familiarity","exposure","proficiency","minimum","maximum","ideally",
    "various","multiple","several","including","following","related","relevant",
    "excellent","strong","good","great","proven","demonstrated","solid",
    "work","working","use","using","used","provide","support","help","ensure",
    "maintain","create","build","deliver","drive","focus","based","given",
    "also","well","both","other","such","like","example","etc","ie","eg",
}


# ============================================================
# HYBRID ANALYSIS ENGINE
# ============================================================

def extract_active_skill_groups(job_text):
    """Find which skill groups are actually required by this job description."""
    job_lower = job_text.lower()
    active = []
    for group in SKILL_GROUPS:
        for pattern in group["jd_patterns"]:
            if re.search(pattern, job_lower):
                active.append(group)
                break
    return active


def match_group_against_cv(group, cv_text):
    """
    For one skill group, check the CV.
    Returns: "strong", "partial", or "missing"
    Also returns evidence (what was found).
    """
    cv_lower = cv_text.lower()
    # Check strong keywords
    found_strong = [kw for kw in group["strong"] if kw in cv_lower]
    if found_strong:
        return "strong", found_strong[0]
    # Check partial keywords
    found_partial = [kw for kw in group["partial"] if kw in cv_lower]
    if found_partial:
        return "partial", found_partial[0]
    return "missing", None


def run_rule_based_analysis(cv_text, job_desc):
    """
    Full rule-based hybrid analysis.
    Returns dict with strong, partial, missing lists + weighted score.
    """
    active_groups = extract_active_skill_groups(job_desc)
    if not active_groups:
        return {"strong": [], "partial": [], "missing": [], "score": 0, "groups": []}

    strong, partial, missing = [], [], []
    total_weight = 0
    earned_weight = 0

    for group in active_groups:
        result, evidence = match_group_against_cv(group, cv_text)
        w = group["weight"]
        total_weight += w
        entry = {"label": group["label"], "evidence": evidence, "category": group["category"], "weight": w}
        if result == "strong":
            strong.append(entry)
            earned_weight += w
        elif result == "partial":
            partial.append(entry)
            earned_weight += w * 0.5   # partial = 50% credit
        else:
            missing.append(entry)

    raw_score = (earned_weight / total_weight * 100) if total_weight > 0 else 0
    # Realistic cap: partial matches penalise slightly, cap at 88
    score = round(min(88, max(20, raw_score)))

    return {
        "strong": strong,
        "partial": partial,
        "missing": missing,
        "score": score,
        "total_reqs": len(active_groups),
        "groups": active_groups,
    }


def get_anthropic_api_key():
    """Return the Anthropic API key from st.secrets, or None if not configured."""
    try:
        return st.secrets["ANTHROPIC_API_KEY"]
    except (KeyError, AttributeError, FileNotFoundError):
        return None


# ============================================================
# CORE AI PIPELINE — 11-STEP CV TAILORING ENGINE
# ============================================================

def claude_api_call(prompt, max_tokens=3000):
    """Central Claude API caller. Returns raw text or None on any failure."""
    api_key = get_anthropic_api_key()
    if not api_key:
        return None
    import urllib.request
    try:
        payload = json.dumps({
            "model": "claude-sonnet-4-6",
            "max_tokens": max_tokens,
            "messages": [{"role": "user", "content": prompt}]
        }).encode()
        req = urllib.request.Request(
            "https://api.anthropic.com/v1/messages",
            data=payload,
            headers={
                "Content-Type": "application/json",
                "anthropic-version": "2023-06-01",
                "x-api-key": api_key,
            }
        )
        with urllib.request.urlopen(req, timeout=60) as resp:
            data = json.loads(resp.read())
            return data["content"][0]["text"]
    except Exception:
        return None


def parse_json_response(text):
    """Strip code fences and parse JSON from a Claude response."""
    if not text:
        return None
    text = re.sub(r'```(?:json)?\s*', '', text).strip().strip('`').strip()
    try:
        return json.loads(text)
    except Exception:
        m = re.search(r'\{[\s\S]*\}', text)
        if m:
            try:
                return json.loads(m.group())
            except Exception:
                pass
    return None


def build_sections_text(cv_sections, max_per_section=700):
    """Format CV sections as labeled text blocks for Claude prompts."""
    out = ""
    order = ["personal_details", "summary", "education", "experience",
             "projects", "skills", "volunteering", "courses_training", "languages"]
    for key in order:
        val = cv_sections.get(key, "").strip()
        if val:
            label = key.upper().replace("_", " ")
            out += f"\n[{label}]\n{val[:max_per_section]}\n"
    return out


# ── STEP 1+2: Deep CV Analysis + Job Requirement Extraction ──
def call_claude_deep_analysis(cv_sections, job_role, job_desc):
    """
    Steps 1+2: Analyze CV structure, identify strengths/weaknesses,
    extract job requirements, compute ATS gaps and match score.
    Returns structured dict or None.
    """
    sections_text = build_sections_text(cv_sections)
    prompt = f"""You are a senior technical recruiter with 15 years of experience. Analyze this CV against the job description and return a precise, honest assessment.

TARGET ROLE: {job_role}

JOB DESCRIPTION:
{job_desc[:2000]}

CV CONTENT:
{sections_text[:3000]}

Return ONLY a JSON object (no markdown, no explanation):
{{
  "cv_strengths": ["specific strength referencing actual CV content", "another strength"],
  "cv_weaknesses": ["specific gap or weakness", "another gap"],
  "hard_skills_required": ["Python", "SQL", "exact hard skill from JD"],
  "soft_skills_required": ["communication", "exact soft skill from JD"],
  "domain_knowledge_required": ["data analysis", "specific domain expertise from JD"],
  "seniority_level": "junior",
  "ats_keywords_missing": ["exact phrase from JD absent from CV", "another keyword"],
  "ats_keywords_present": ["keyword found in both JD and CV"],
  "match_score": 68,
  "match_label": "Moderate Match",
  "score_rationale": "Honest 2-sentence explanation referencing actual CV content and JD requirements.",
  "metrics_missing": ["Experience bullet that needs a number or outcome added"],
  "quick_wins": ["Most impactful improvement #1", "Most impactful improvement #2", "Most impactful improvement #3"]
}}

STRICT RULES:
- match_score must be 0-100 integer — be realistic, not generous
- ats_keywords_missing: copy exact phrases/words from JD that are absent from the CV
- metrics_missing: paraphrase actual bullets from the CV that would benefit from quantification
- cv_strengths: cite only what is actually in this CV — no generic praise
- match_label must be one of: "Weak Match", "Moderate Match", "Strong Match" """

    raw = claude_api_call(prompt, max_tokens=2000)
    return parse_json_response(raw)


# ── STEP 3+4: Dynamic Question Generation & Structured Memory ──
def call_claude_generate_questions(cv_sections, job_role, job_desc, ai_analysis):
    """
    Step 3: Generate 5-8 specific, personalized questions targeting gaps
    identified in the deep analysis. NOT generic questions.
    Returns list of question dicts or [].
    """
    sections_text = build_sections_text(cv_sections, max_per_section=400)
    ats_missing    = (ai_analysis or {}).get("ats_keywords_missing", [])[:8]
    weaknesses     = (ai_analysis or {}).get("cv_weaknesses", [])[:5]
    metrics_gaps   = (ai_analysis or {}).get("metrics_missing", [])[:4]

    prompt = f"""You are a career coach helping a candidate strengthen their CV application. Generate 5-8 highly specific, personalized questions to uncover information that will make this CV noticeably stronger.

TARGET ROLE: {job_role}

CV GAPS IDENTIFIED BY ANALYSIS:
- Missing ATS keywords: {', '.join(ats_missing) or 'see CV below'}
- CV weaknesses: {', '.join(weaknesses) or 'see CV below'}
- Bullets needing metrics: {', '.join(metrics_gaps) or 'see CV below'}

CURRENT CV (abbreviated):
{sections_text[:2500]}

QUESTION QUALITY BENCHMARK:
BAD: "Do you have SQL experience?" — too closed, not personalized
GOOD: "Your CV shows Python for data analysis — have you also used SQL to query databases, even in university coursework or side projects?"

BAD: "Tell us about your leadership skills"
GOOD: "You mention tutoring students — how many did you tutor, over what time period, and did you notice measurable improvement in their academic performance?"

BAD: "Describe a challenge you overcame"
GOOD: "Your network analysis project — what was the dataset size, and what specific finding or result did the analysis produce?"

Generate questions that:
1. Reference SPECIFIC content actually in this CV (name the project, role, or skill)
2. Target the identified gaps and missing keywords from the job description
3. Uncover quantifiable achievements (scope, numbers, outcomes)
4. Surface relevant experience with JD keywords the CV is missing

Return ONLY JSON:
{{
  "questions": [
    {{
      "key": "q_unique_snake_case_key",
      "question": "Full specific question text referencing their actual CV",
      "hint": "e.g., 'I tutored 15 students over 6 months, their grades improved by...'",
      "section": "experience",
      "gap_addressed": "Adds scope and outcome metrics to tutoring entry"
    }}
  ]
}}

Generate exactly 5-8 questions. Every question must be specific to THIS person's CV."""

    raw = claude_api_call(prompt, max_tokens=1500)
    result = parse_json_response(raw)
    if result and isinstance(result.get("questions"), list):
        return result["questions"]
    return []


# ── STEPS 5-10: Full CV Rewrite — ATS, Summary, Bullets, Localization Prep ──
def call_claude_rewrite_cv(cv_sections, job_role, job_desc, ai_analysis, follow_up_answers):
    """
    Steps 5-10: True CV reconstruction using original CV + job description
    + all user answers. Produces professional-recruiter-quality output.
    Returns dict with rewritten sections + improvement_log, or None.
    """
    sections_text = build_sections_text(cv_sections)

    # Format answers preserving the structured memory (Step 4)
    answers_formatted = ""
    if follow_up_answers:
        for key, val in follow_up_answers.items():
            if isinstance(val, dict) and val.get("answer", "").strip():
                q = val.get("question", key)
                a = val["answer"].strip()
                answers_formatted += f"Q: {q}\nA: {a}\n\n"
            elif isinstance(val, str) and val.strip():
                answers_formatted += f"- {val.strip()}\n"

    ats_missing  = (ai_analysis or {}).get("ats_keywords_missing", [])[:12]
    strengths    = (ai_analysis or {}).get("cv_strengths", [])[:5]
    hard_skills  = (ai_analysis or {}).get("hard_skills_required", [])[:8]

    prompt = f"""You are a professional CV writer with 15 years of experience placing candidates at top companies. Rewrite this CV to be a compelling, targeted application. Think like a human recruiter reading hundreds of CVs — make this one stand out.

TARGET ROLE: {job_role}

JOB DESCRIPTION:
{job_desc[:1500]}

ANALYSIS RESULTS:
- ATS keywords to inject naturally (only where genuinely applicable): {', '.join(ats_missing) or 'none'}
- Confirmed strengths to emphasize: {', '.join(strengths) or 'see CV'}
- Hard skills required by the role: {', '.join(hard_skills) or 'see JD'}

USER'S ANSWERS (CONFIRMED REAL INFORMATION — integrate these actively into the rewrite):
{answers_formatted if answers_formatted else "No additional answers provided — rewrite from CV content only."}

ORIGINAL CV:
{sections_text}

=== YOUR REWRITING INSTRUCTIONS ===

[PROFESSIONAL SUMMARY] — Write a completely new 3-4 sentence paragraph:
  • Open with the candidate's level + target role: "Final-year [field] student / [X]-year professional targeting a [role] position..."
  • Name their 2-3 strongest relevant technical skills with specific tools mentioned in the CV
  • Include their top achievement or differentiator (from CV or user answers)
  • Close with what value they bring — confident, active voice
  • AVOID: "hardworking", "passionate learner", "team player" as empty openers

[EXPERIENCE] — For EACH position, rewrite bullet points to:
  • Begin with a strong past-tense action verb: Developed, Engineered, Analyzed, Led, Designed, Optimized, Delivered, Streamlined, Reduced, Increased, Automated, Built
  • Add specifics from user answers (numbers of people, duration, outcome) if provided
  • Inject 1-2 JD keywords per bullet where they naturally apply — never force them
  • One bullet = one achievement, not a task description

[PROJECTS] — Rewrite each project entry to:
  • Lead with what was built and what tool: "Built a [X] using [Y] to [accomplish Z]"
  • Include metrics (dataset rows, accuracy, users, time saved) IF user mentioned them in answers
  • End with the impact or insight produced

[SKILLS] — Reorder: most JD-relevant skills first. Add any skills the user confirmed in answers.

[EDUCATION] — Keep all dates and institution names EXACTLY. Minor wording improvements only.

[VOLUNTEERING/ACTIVITIES] — Strengthen action verbs. Keep facts identical.

=== HARD RULES (violation = the output is wrong) ===
• NEVER add companies, degrees, job titles, or dates not in the original CV
• NEVER invent metrics (numbers, %, timeframes) unless the user explicitly stated them in their answers above
• If user said "I don't have X" — do NOT add X anywhere
• All proper nouns (university names, company names, tool names, city names) stay exactly as written
• Only include sections that exist in the original CV

Return ONLY valid JSON (no markdown fences, no preamble, no explanation):
{{
  "summary": "new professional summary paragraph",
  "experience": "full rewritten experience section — keep all entries, strengthen all bullets",
  "projects": "full rewritten projects section",
  "skills": "reordered, enriched skills list",
  "education": "education — minor improvements only, all names/dates identical",
  "volunteering": "volunteering section",
  "courses_training": "courses — unchanged",
  "languages": "languages — unchanged",
  "improvement_log": [
    {{"type": "Improved", "section": "Summary", "description": "Rewrote to target {job_role} explicitly, led with Python/SQL as primary skills"}},
    {{"type": "Added", "section": "Experience", "description": "Incorporated scope from user answer: tutored 40+ students over 2 semesters"}},
    {{"type": "Reorganized", "section": "Skills", "description": "Moved Python and SQL to top — both are core requirements in JD"}}
  ]
}}"""

    raw = claude_api_call(prompt, max_tokens=4500)
    return parse_json_response(raw)


# ── STEP 9: Professional Multilingual Localization via Claude ──
def call_claude_localize_cv(cv_data, target_lang):
    """
    Step 9: Localize the CV to Hebrew or Arabic using Claude.
    Native-level professional phrasing — not word-by-word translation.
    Falls back to googletrans if Claude unavailable.
    """
    lang_style = {
        "Hebrew": (
            "professional Israeli Hebrew (עברית מקצועית לשוק העבודה הישראלי). "
            "Use formal Israeli CV conventions and natural Hebrew phrasing — "
            "this is LOCALIZATION by a native Israeli CV writer, not a word-by-word translation."
        ),
        "Arabic": (
            "professional Modern Standard Arabic suitable for the Israeli job market. "
            "Use formal Arabic CV conventions — "
            "this is LOCALIZATION by a native Arabic CV writer, not a word-by-word translation."
        ),
    }

    sections_to_translate = {}
    for key in ["summary", "education", "experience", "projects",
                "skills", "volunteering", "courses_training", "languages"]:
        val = cv_data.get(key, "").strip()
        if val:
            sections_to_translate[key] = val

    if not sections_to_translate:
        return cv_data

    sections_json = json.dumps(sections_to_translate, ensure_ascii=False)

    prompt = f"""Translate and localize this CV content to {lang_style[target_lang]}.

MANDATORY RULES:
- Keep ALL technical tool names in English: Python, SQL, Excel, GitHub, Tableau, TensorFlow, etc.
- Keep ALL institution and company names in their original language
- Keep ALL dates EXACTLY as written (2022–2025, Present, etc.)
- Keep ALL numbers, percentages, and GPA values exactly
- For Hebrew: use masculine formal forms as default
- Do NOT translate proper nouns (universities, companies, cities, person names)

CV SECTIONS TO LOCALIZE (JSON):
{sections_json}

Return the same JSON keys with localized text. Return ONLY the JSON object — no explanation, no markdown."""

    raw = claude_api_call(prompt, max_tokens=4000)
    result = parse_json_response(raw)

    if result and isinstance(result, dict):
        final = cv_data.copy()
        for key, val in result.items():
            if val and isinstance(val, str) and val.strip():
                final[key] = val
        return final

    # Fallback to googletrans if Claude unavailable or failed
    return translate_cv_googletrans(cv_data, target_lang)

GENERIC_IGNORE_WORDS = NOISE_WORDS  # alias for backward compat

ACTION_VERBS = {
    "assisted": "Engineered","helped": "Optimized","did": "Developed",
    "worked on": "Spearheaded","responsible for": "Directed","learned": "Mastered",
    "managed": "Architected","wrote": "Formulated"
}

HEADINGS_MAP = {
    "English": ["Personal Details","Professional Summary","Education","Professional Experience",
                "Projects","Courses & Training","Volunteering / Community","Languages","Skills","Additional Information"],
    "Hebrew":  ["פרטים אישיים","תמצית מקצועית","השכלה","ניסיון תעסוקתי","פרויקטים",
                "קורסים והכשרות","התנדבות / קהילה","שפות","כישורים","מידע נוסף"],
    "Arabic":  ["البيانات الشخصية","الملخص المهني","التعليم","الخبرة العملية","المشاريع",
                "الدورات والتدريب","التطوع / المجتمع","اللغات","المهارات","معلومات إضافية"]
}

SECTION_REGEX = {
    "personal_details": r"(personal details|contact|contact info|פרטים אישיים|البيانات الشخصية)",
    "summary": r"(summary|profile|about me|professional summary|תמצית|פרופיל|תמצית מקצועית|الملخص|ملخص مهني)",
    "education": r"(education|academic|degrees|השכלה|לימודים|التعليم|المؤهلات العلمية)",
    "experience": r"(experience|work|employment|history|professional experience|ניסיון|ניסיון תעסוקתי|الخبرة|الخبرة العملية)",
    "projects": r"(projects|key projects|academic projects|פרויקטים|פרוייקטים|المشاريع)",
    "courses_training": r"(courses|training|certification|certifications|קורסים|הכשרות|הסמכות|الدورات|التدريب)",
    "volunteering": r"(volunteering|community|extracurricular|volunteer|leadership|activities|התנדבות|פעילות קהילתית|التطوع)",
    "languages": r"(languages|שפות|اللغات)",
    "skills": r"(skills|competencies|technologies|technical skills|כישורים|מיומנויות|יכולות|المهارات)"
}


# ============================================================
# CORE FUNCTIONS
# ============================================================

def detect_language(text):
    if re.search(r'[\u0590-\u05FF]', text): return "Hebrew"
    if re.search(r'[\u0600-\u06FF]', text): return "Arabic"
    return "English"


def validate_job_description(job_text):
    words = job_text.strip().split()
    if len(words) < 30:
        return False, "Please paste a real job description (at least a few sentences with requirements)."
    unique_words = {w.lower() for w in words if len(w) > 2}
    if len(unique_words) < 10:
        return False, "Please paste a real job description with responsibilities and requirements."
    return True, ""


def extract_text_from_pdf(file):
    try:
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    except:
        return ""


def extract_text_from_docx(file):
    try:
        doc = docx.Document(file)
        return "\n".join([para.text for para in doc.paragraphs])
    except:
        return ""


def classify_sections_refined(text):
    sections = {
        "personal_details":"","summary":"","education":"","experience":"",
        "projects":"","courses_training":"","volunteering":"","languages":"",
        "skills":"","additional_information":""
    }
    # Insert newlines before known section headers (handles flat PDF)
    ALL_CAPS_HEADERS = [
        "PROFESSIONAL SUMMARY","PROFESSIONAL EXPERIENCE","WORK EXPERIENCE",
        "EDUCATION","PROJECTS","TECHNICAL SKILLS","SKILLS","LANGUAGES",
        "COURSES","VOLUNTEERING","LEADERSHIP & ACTIVITIES","LEADERSHIP",
        "ACTIVITIES","CERTIFICATIONS","TRAINING","ACADEMIC BACKGROUND",
        "תמצית מקצועית","השכלה","ניסיון תעסוקתי","פרויקטים","כישורים","שפות","התנדבות",
        "الملخص المهني","التعليم","الخبرة العملية","المشاريع","المهارات","اللغات","التطوع",
    ]
    for hw in sorted(ALL_CAPS_HEADERS, key=len, reverse=True):
        text = text.replace(hw, f"\n\n{hw}\n")
    current_section = "personal_details"
    lines = text.split('\n')
    for line in lines:
        clean_line = line.strip()
        if not clean_line: continue
        low_line = clean_line.lower()
        found_header = False
        if len(clean_line) < 60:
            for key, pattern in SECTION_REGEX.items():
                # Require the pattern to be the whole line (allow leading/trailing
                # decoration like dashes, colons, whitespace — but no other words).
                if re.search(r'^[\s\-=*•·_|]*' + pattern + r'[\s\-=*•·_|:]*$', low_line, re.IGNORECASE):
                    current_section = key
                    found_header = True
                    break
        if not found_header:
            sections[current_section] += clean_line + "\n"
    # Trim personal_details to just name + contact (max 5 lines)
    pd_lines = [l for l in sections["personal_details"].split('\n') if l.strip()]
    if len(pd_lines) > 5:
        sections["personal_details"] = "\n".join(pd_lines[:5])
    for k in sections:
        sections[k] = re.sub(r'\n{3,}', '\n\n', sections[k].strip())
    return sections


def format_contact_info(text):
    if not text: return "CANDIDATE", ""
    # Normalize separators
    text_clean = re.sub(r'\s*[|·•]\s*', '\n', text)
    lines = [l.strip() for l in text_clean.split('\n') if l.strip()]
    name, contact_parts = "", []
    for i, line in enumerate(lines):
        low = line.lower()
        is_contact = (any(k in low for k in ['email','phone','@','linkedin','github','israel','jerusalem','city','tel aviv'])
                      or re.search(r'\d{3}', line) or '@' in line)
        if i == 0 and not is_contact and len(line.split()) <= 6:
            name = line
        else:
            contact_parts.append(line)
    if not name and lines:
        m = re.match(r'^([A-Za-z\u0590-\u05FF\u0600-\u06FF ]{2,40}?)(?=\s+\S*[@\d])', text)
        name = m.group(1).strip() if m else lines[0]
        if not contact_parts: contact_parts = lines[1:]
    contact_str = " | ".join(contact_parts)
    contact_str = re.sub(r'\b(Email|Phone|Tel|Mobile|LinkedIn|Address)\s*:\s*', '', contact_str, flags=re.IGNORECASE)
    return name.strip(), contact_str.strip()


def group_skills_segregated(skills_text):
    if not skills_text: return {"PROFESSIONAL": [], "TECHNICAL": []}
    # Remove subheading labels that bleed in from PDF
    cleaned = re.sub(r'(Technical Skills|Business & Professional Skills|Professional Skills)\s*[,\n]?', '', skills_text, flags=re.IGNORECASE)
    skills = [s.strip() for s in re.split(r'[,\n•|]', cleaned) if s.strip() and len(s.strip()) > 1]
    tech_kw = set(list(SKILL_KEYWORDS) + ["sql","python","aws","docker","git","react","tableau","powerbi",
                                           "excel","data analysis","r","networkx","pandas","numpy","scipy",
                                           "matplotlib","unix","bash","bigquery","spss","statistical modeling",
                                           "machine learning","network analysis","rstudio","bigquery"])
    technical, professional = [], []
    for s in skills:
        sl = s.lower()
        if any(kw in sl for kw in tech_kw):
            technical.append(s)
        else:
            professional.append(s)
    return {"PROFESSIONAL": list(dict.fromkeys(professional)), "TECHNICAL": list(dict.fromkeys(technical))}


def build_final_cv(cv_sections, adjusted_sections, job_role):
    """
    Merges adjusted sections back into cv_sections.
    Falls back to original section if adjustment is missing or empty.
    Always preserves: personal_details exactly as-is.
    """
    final = cv_sections.copy()
    if adjusted_sections:
        for key, val in adjusted_sections.items():
            if val and val.strip() and key != "personal_details":
                final[key] = val.strip()
    # Always keep personal_details untouched
    final["personal_details"] = cv_sections.get("personal_details", "")
    return final


def split_cv_entries(text):
    """Split flat CV text into separate entries, handling inline • bullets."""
    if not text: return []
    # Expand inline bullets to newlines
    text = re.sub(r'\s*•\s*', '\n• ', text)
    # If double newlines exist, use them
    if '\n\n' in text:
        return [b.strip() for b in text.split('\n\n') if b.strip()]
    # Otherwise split: new entry starts when a non-bullet line follows bullet lines
    lines = text.split('\n')
    blocks, current = [], []
    for line in lines:
        stripped = line.strip()
        if not stripped: continue
        is_bullet = stripped.startswith(('•','-','*','·'))
        if current and not is_bullet and not current[-1].startswith('•'):
            blocks.append('\n'.join(current))
            current = [stripped]
        else:
            current.append(stripped)
    if current: blocks.append('\n'.join(current))
    return [b for b in blocks if b.strip()] or [text]


def render_entry(block, show_subtitle=True):
    """Render one CV entry block → HTML."""
    block = re.sub(r'\s*•\s*', '\n• ', block)
    lines = [l.strip() for l in block.split('\n') if l.strip()]
    if not lines: return ""
    title = lines[0].lstrip('•-* ')
    html_out = f'<div class="cv-entry"><div class="cv-entry-title">{title}</div>'
    subtitle_done = False
    bullets = []
    for line in lines[1:]:
        is_bullet = line.startswith(('•','-','*'))
        clean = line.lstrip('•-* ').strip()
        if not clean: continue
        if is_bullet:
            bullets.append(clean)
        elif not subtitle_done and show_subtitle:
            html_out += f'<div class="cv-entry-sub">{clean}</div>'
            subtitle_done = True
        else:
            bullets.append(clean)
    if bullets:
        html_out += '<ul class="cv-bullets">' + ''.join(f'<li>{b}</li>' for b in bullets) + '</ul>'
    html_out += '</div>'
    return html_out


def translate_cv_googletrans(cv_data, target_lang):
    """
    Translates CV sections using deep-translator.
    Uses word-boundary-safe placeholders to avoid bugs like comMUNication.
    """
    try:
        from deep_translator import GoogleTranslator
    except ImportError:
        return cv_data

    lang_code = "iw" if target_lang == "Hebrew" else "ar"

    # Ordered longest-first to avoid partial matches
    PRESERVE = [
        "Hebrew University of Jerusalem",
        "Red Sea Candles",
        "Rosary Sisters High School",
        "Nayzak Program",
        "HUBS Aid Program",
        "Model United Nations",
        "scikit-learn", "RStudio", "BigQuery", "NetworkX",
        "NumPy", "matplotlib", "scipy", "pandas",
        "TensorFlow", "PyTorch", "GitHub", "GitLab",
        "Streamlit", "Jupyter", "Mechina", "Tawjihi", "IGCSE",
        "Python", "Unix", "Bash", "Linux", "Docker",
        "Excel", "Tableau", "Azure", "HUJI", "HUBS",
        "Git", "AWS", "GPA", "MVP", "CLI",
        "Facebook", "LinkedIn",
    ]

    def protect_and_translate(text):
        if not text or not text.strip():
            return text

        placeholders = {}
        protected = text
        counter = 0

        # Step 1: protect multi-word phrases and tools using word boundaries
        for word in PRESERVE:
            # Use word boundary only for short single words, not phrases
            if ' ' in word:
                pattern = re.escape(word)
            else:
                pattern = r'\b' + re.escape(word) + r'\b'
            ph = f"KEEP{counter}KEEP"
            new_protected = re.sub(pattern, ph, protected, flags=re.IGNORECASE)
            if new_protected != protected:
                placeholders[ph] = word
                protected = new_protected
                counter += 1

        # Step 2: protect dates like "2025–Present", "2025-2026", "2025"
        for m in re.finditer(r'\d{4}[–\-](?:Present|\d{4})|\b\d{4}\b', protected):
            ph = f"DATE{counter}DATE"
            placeholders[ph] = m.group()
            protected = protected[:m.start()] + ph + protected[m.end():]
            counter += 1
            # Restart after modification to avoid index issues
            break
        # Simpler date protection — replace all at once
        def replace_date(m):
            nonlocal counter
            ph = f"DATE{counter}DATE"
            placeholders[ph] = m.group()
            counter += 1
            return ph
        protected = re.sub(r'\d{4}[–\-](?:Present|\d{4})|\b\d{4}\b', replace_date, protected)

        # Step 3: protect score numbers like "90.4", "96"
        protected = re.sub(r'\b\d+\.\d+\b|\b\d{2,3}\b',
                           lambda m: (placeholders.__setitem__(f"NUM{counter}NUM", m.group()) or f"NUM{counter}NUM"),
                           protected)

        # Step 4: translate line by line
        lines = protected.split('\n')
        translated_lines = []
        translator = GoogleTranslator(source='auto', target=lang_code)

        for line in lines:
            stripped = line.strip()
            if not stripped:
                translated_lines.append('')
                continue
            # Skip if line has no real content after removing placeholders
            content_only = re.sub(r'\w+\d+\w+', '', stripped).strip('•-* \t|·–—:,.')
            if not content_only:
                translated_lines.append(line)
                continue
            try:
                if len(stripped) > 4500:
                    translated_lines.append(line)
                    continue
                t = translator.translate(stripped)
                translated_lines.append(t if t and t.strip() else line)
            except Exception:
                translated_lines.append(line)

        result = '\n'.join(translated_lines)

        # Step 5: restore all placeholders
        for ph, original in placeholders.items():
            result = result.replace(ph, original)

        return result

    result = cv_data.copy()
    for key in ["summary", "education", "experience", "projects",
                "courses_training", "volunteering", "languages", "skills"]:
        val = cv_data.get(key, "").strip()
        if not val:
            continue
        try:
            translated = protect_and_translate(val)
            if translated and translated.strip():
                result[key] = translated
        except Exception:
            pass

    return result


def extract_name_and_contact(personal_details_text):
    """
    Robustly extracts name and contact line from personal_details.
    Handles flat PDF text, structured text, and pipe-separated formats.
    Returns (name, contact_line) — name is NEVER repeated in contact_line.
    """
    text = personal_details_text.strip()
    if not text:
        return "CANDIDATE", ""

    CONTACT_KEYWORDS = ['email', 'phone', 'tel', 'mobile', '@', 'linkedin',
                        'github', 'israel', 'jerusalem', 'city', 'address',
                        'ירושלים', 'תל אביב', 'القدس', 'إسرائيل']

    # ── Strategy 1: newline-separated (clean PDF/manual) ──
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    if len(lines) >= 2:
        first = lines[0]
        low = first.lower()
        has_contact = any(k in low for k in CONTACT_KEYWORDS) or '@' in first
        is_name_length = len(first.split()) <= 5 and len(first) <= 45
        if not has_contact and is_name_length:
            name = first
            # Contact = everything after first line, joined
            rest = " | ".join(lines[1:])
            rest = re.sub(r'\b(Email|Phone|Tel|Mobile|E-mail)\s*:\s*', '', rest, flags=re.IGNORECASE)
            # Remove name if it appears at start of contact line
            if rest.upper().startswith(name.upper()):
                rest = rest[len(name):].strip(' |·,')
            contact = rest.strip(' |·')
            return name.strip(), contact

    # ── Strategy 2: flat single line — extract name from beginning ──
    # Pattern: "FIRSTNAME LASTNAME City, Country Email: ... Phone: ..."
    flat = text.replace('\n', ' ')

    # Try: name = first all-caps word(s) before a city/country keyword
    m = re.match(
        r'^([A-Z][A-Z\s]{2,30}?)(?=\s+(?:Jerusalem|Israel|Tel Aviv|Haifa|Email|Phone|\d{3}|ירושלים|القدس))',
        flat
    )
    if m:
        name = m.group(1).strip()
        rest = flat[len(name):].strip()
        rest = re.sub(r'\b(Email|Phone|Tel|Mobile|E-mail)\s*:\s*', '', rest, flags=re.IGNORECASE)
        # Remove leading duplicate of name
        if rest.upper().startswith(name.upper()):
            rest = rest[len(name):].strip(' |·,')
        # Build clean contact from rest
        parts = re.split(r'\s*[|·•]\s*', rest)
        parts = [p.strip() for p in parts if p.strip() and p.strip().upper() != name.upper()]
        return name, " | ".join(parts)

    # ── Strategy 3: pipe-separated "Name | City | Email | Phone" ──
    if '|' in flat:
        parts = [p.strip() for p in flat.split('|')]
        if parts:
            name = parts[0]
            contact = " | ".join(parts[1:])
            contact = re.sub(r'\b(Email|Phone|Tel|Mobile)\s*:\s*', '', contact, flags=re.IGNORECASE)
            return name.strip(), contact.strip()

    # ── Strategy 4: last resort — first 2-3 words as name ──
    words = flat.split()
    name = ' '.join(words[:3]) if len(words) >= 3 else flat
    rest = ' '.join(words[3:]) if len(words) > 3 else ''
    rest = re.sub(r'\b(Email|Phone|Tel|Mobile)\s*:\s*', '', rest, flags=re.IGNORECASE)
    return name.strip(), rest.strip()


def render_cv_html(cv_data, lang="English"):
    """
    Renders CV as clean Israeli-professional HTML.
    Section order: Summary → Education → Technical Skills → Projects → Experience → Volunteering → Languages
    """
    is_rtl = lang in ["Hebrew", "Arabic"]
    dir_attr = 'dir="rtl"' if is_rtl else 'dir="ltr"'
    text_align = 'right' if is_rtl else 'left'
    h = HEADINGS_MAP[lang]

    name, contact = extract_name_and_contact(cv_data.get("personal_details", ""))
    if not name: name = "CANDIDATE"

    # ── CSS embedded in the CV page ──
    cv_css = f"""
    <style>
    .isr-cv {{
        background: white;
        max-width: 780px;
        margin: 0 auto;
        padding: 40px 50px;
        font-family: 'Calibri', 'Arial', sans-serif;
        font-size: 10.5pt;
        color: #111;
        line-height: 1.45;
        direction: {('rtl' if is_rtl else 'ltr')};
        text-align: {text_align};
        box-shadow: 0 4px 24px rgba(0,0,0,0.10);
    }}
    .isr-name {{
        font-size: 20pt;
        font-weight: 700;
        text-align: center;
        letter-spacing: 0.5px;
        margin-bottom: 4px;
        color: #0f172a;
        text-transform: uppercase;
    }}
    .isr-contact {{
        font-size: 10pt;
        text-align: center;
        color: #334155;
        margin-bottom: 10px;
    }}
    .isr-divider {{
        border: none;
        border-top: 1.5px solid #0f172a;
        margin: 8px 0 14px 0;
    }}
    .isr-section-title {{
        font-size: 11.5pt;
        font-weight: 700;
        color: #1e293b;
        text-transform: uppercase;
        letter-spacing: 0.8px;
        border-bottom: 1px solid #94a3b8;
        padding-bottom: 2px;
        margin-top: 14px;
        margin-bottom: 6px;
    }}
    .isr-entry {{ margin-bottom: 8px; }}
    .isr-entry-header {{
        display: flex;
        justify-content: space-between;
        align-items: baseline;
        flex-direction: {'row-reverse' if is_rtl else 'row'};
    }}
    .isr-entry-title {{ font-weight: 700; font-size: 10.5pt; color: #0f172a; }}
    .isr-entry-date  {{ font-size: 9.5pt; color: #475569; font-weight: 600; white-space: nowrap; }}
    .isr-entry-sub   {{ font-size: 10pt; color: #475569; font-style: italic; margin-bottom: 3px; }}
    .isr-bullets     {{ margin: 3px 0; {'padding-right:16px;padding-left:0;' if is_rtl else 'padding-left:16px;'} }}
    .isr-bullets li  {{ font-size: 10pt; color: #1e293b; margin-bottom: 2px; }}
    .isr-skills-row  {{ margin-bottom: 4px; font-size: 10pt; }}
    .isr-skills-label{{ font-weight: 700; color: #1e293b; }}
    </style>
    """

    def sec(title):
        return f'<div class="isr-section-title">{title}</div>'

    def parse_date_from_title(title):
        """Extract trailing date from title like 'Company Name | 2025–Present'"""
        m = re.search(r'[\|｜]\s*(\d{4}[–\-]\w+|\d{4})\s*$', title)
        if m:
            date = m.group(1)
            clean_title = title[:m.start()].strip()
            return clean_title, date
        return title, ""

    def render_entries_html(content, show_date=True):
        out = ""
        for block in split_cv_entries(content):
            block = re.sub(r'\s*•\s*', '\n• ', block)
            lines = [l.strip() for l in block.split('\n') if l.strip()]
            if not lines: continue

            raw_title = lines[0].lstrip('•-* ')
            title, date = parse_date_from_title(raw_title)

            out += '<div class="isr-entry">'
            out += '<div class="isr-entry-header">'
            out += f'<span class="isr-entry-title">{title}</span>'
            if date:
                out += f'<span class="isr-entry-date">{date}</span>'
            out += '</div>'

            subtitle_done = False
            bullets = []
            for line in lines[1:]:
                is_bullet = line.startswith(('•','-','*'))
                clean = line.lstrip('•-* ').strip()
                if not clean: continue
                if is_bullet:
                    bullets.append(clean)
                elif not subtitle_done:
                    out += f'<div class="isr-entry-sub">{clean}</div>'
                    subtitle_done = True
                else:
                    bullets.append(clean)

            if bullets:
                out += '<ul class="isr-bullets">' + ''.join(f'<li>{b}</li>' for b in bullets) + '</ul>'
            out += '</div>'
        return out

    # ── BUILD HTML ──
    html = f'{cv_css}<div class="isr-cv" {dir_attr}>'

    # Header
    html += f'<div class="isr-name">{name}</div>'
    if contact:
        html += f'<div class="isr-contact">{contact}</div>'
    html += '<hr class="isr-divider">'

    # 1. SUMMARY
    if cv_data.get("summary"):
        html += sec(h[1])
        summary = cv_data["summary"].replace('\n', ' ').strip()
        html += f'<div style="font-size:10.5pt;color:#1e293b;line-height:1.5;margin-bottom:8px;">{summary}</div>'

    # 2. EDUCATION
    if cv_data.get("education"):
        html += sec(h[2])
        edu = re.sub(r'\s*•\s*', '\n• ', cv_data["education"])
        # Split on lines that start with a capital and look like institution names
        edu_blocks = re.split(r'\n(?=[A-Z\u0590-\u05FF\u0600-\u06FF])', edu)
        for block in edu_blocks:
            lines = [l.strip() for l in block.split('\n') if l.strip()]
            if not lines: continue
            raw_title = lines[0].lstrip('•-* ')
            title, date = parse_date_from_title(raw_title)
            html += '<div class="isr-entry">'
            html += '<div class="isr-entry-header">'
            html += f'<span class="isr-entry-title">{title}</span>'
            if date: html += f'<span class="isr-entry-date">{date}</span>'
            html += '</div>'
            for l in lines[1:]:
                clean = l.lstrip('•-* ').strip()
                if clean:
                    html += f'<div class="isr-entry-sub">{clean}</div>'
            html += '</div>'

    # 3. EXPERIENCE
    if cv_data.get("experience"):
        html += sec(h[3])
        html += render_entries_html(cv_data["experience"])

    # 4. PROJECTS
    if cv_data.get("projects"):
        html += sec(h[4])
        html += render_entries_html(cv_data["projects"])

    # 5. SKILLS
    raw_skills = cv_data.get("skills", "")
    if raw_skills:
        raw_skills_clean = re.sub(
            r'(Technical Skills|Business & Professional Skills|Professional Skills|'
            r'כישורים טכניים|כישורים מקצועיים|المهارات التقنية|المهارات المهنية)\s*[,\n]?',
            '', raw_skills, flags=re.IGNORECASE)
        skills_data = group_skills_segregated(raw_skills_clean)

        # Skills section title
        skills_title = h[8] if len(h) > 8 else ("כישורים" if lang == "Hebrew" else "المهارات" if lang == "Arabic" else "Skills")
        html += sec(skills_title)

        if skills_data["TECHNICAL"]:
            tech_label = ("תכנות וכלים" if lang == "Hebrew" else "البرمجة والأدوات" if lang == "Arabic" else "Technical")
            html += f'<div class="isr-skills-row"><span class="isr-skills-label">{tech_label}:</span> {" · ".join(skills_data["TECHNICAL"])}</div>'
        if skills_data["PROFESSIONAL"]:
            prof_label = ("כישורים מקצועיים" if lang == "Hebrew" else "المهارات المهنية" if lang == "Arabic" else "Professional")
            html += f'<div class="isr-skills-row"><span class="isr-skills-label">{prof_label}:</span> {" · ".join(skills_data["PROFESSIONAL"])}</div>'

    # 6. VOLUNTEERING / AWARDS / LEADERSHIP
    if cv_data.get("volunteering"):
        html += sec(h[6])
        vol = re.sub(r'^\s*[&\*]\s*$', '', cv_data["volunteering"], flags=re.MULTILINE)
        html += render_entries_html(vol)

    # 7. LANGUAGES
    if cv_data.get("languages"):
        html += sec(h[7])
        langs = re.sub(r'\s*•\s*', ' · ', cv_data["languages"]).strip(' ·')
        html += f'<div style="font-size:10pt;color:#1e293b;">{langs}</div>'

    # 8. COURSES (if exists)
    if cv_data.get("courses_training"):
        html += sec(h[5])
        items = [c.strip() for c in re.split(r'[,\n]', cv_data["courses_training"]) if c.strip() and len(c.strip()) > 3]
        html += '<ul class="isr-bullets">' + ''.join(f'<li>{i}</li>' for i in items) + '</ul>'

    html += '</div>'
    return html


def generate_docx(cv_data, output_lang):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    for sec in doc.sections:
        sec.top_margin  = docx.shared.Cm(1.8)
        sec.bottom_margin = docx.shared.Cm(1.8)
        sec.left_margin  = docx.shared.Cm(2.0)
        sec.right_margin = docx.shared.Cm(2.0)

    is_rtl = output_lang in ["Hebrew","Arabic"]
    h = HEADINGS_MAP[output_lang]
    font_name = 'Arial' if is_rtl else 'Calibri'

    style = doc.styles['Normal']
    style.font.name = font_name
    style.font.size = Pt(10.5)

    def set_para_format(para, rtl=False):
        if rtl:
            pPr = para._p.get_or_add_pPr()
            bidi = OxmlElement('w:bidi')
            pPr.append(bidi)
            jc = OxmlElement('w:jc')
            jc.set(qn('w:val'), 'right')
            pPr.append(jc)

    def add_para(text, bold=False, italic=False, size=10.5, align="left", color_rgb=None, space_before=0, space_after=2):
        p = doc.add_paragraph()
        if align == "center": p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "right" or is_rtl: p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else: p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_para_format(p, is_rtl)
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        run = p.add_run(str(text))
        run.bold = bold; run.italic = italic
        run.font.name = font_name
        run.font.size = Pt(size)
        if color_rgb:
            run.font.color.rgb = RGBColor(*color_rgb)
        return p

    def add_section_heading(title):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if is_rtl else WD_ALIGN_PARAGRAPH.LEFT
        set_para_format(p, is_rtl)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run(title.upper())
        run.bold = True; run.font.size = Pt(11)
        run.font.name = font_name
        run.font.color.rgb = RGBColor(30, 41, 59)
        # Bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '94a3b8')
        pBdr.append(bottom)
        pPr.append(pBdr)

    def add_bullet_item(text):
        p = doc.add_paragraph(style='List Bullet')
        set_para_format(p, is_rtl)
        if is_rtl: p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(text.lstrip('•-* ').strip())
        run.font.size = Pt(10); run.font.name = font_name
        p.paragraph_format.space_after = Pt(1)

    def parse_date_from_title(title):
        m = re.search(r'[\|｜]\s*(\d{4}[–\-]\w+|\d{4})\s*$', title)
        if m:
            return title[:m.start()].strip(), m.group(1)
        return title, ""

    # ── NAME ──
    name, contact = extract_name_and_contact(cv_data.get("personal_details",""))
    if not name: name = "CANDIDATE"
    name_p = add_para(name.upper(), bold=True, size=18, align="center", space_after=2)
    if contact:
        contact_clean = re.sub(r'\b(Email|Phone|Tel)\s*:\s*', '', contact, flags=re.IGNORECASE)
        add_para(contact_clean, size=10, align="center", space_after=4)

    # Divider line (using a paragraph with border)
    div_p = doc.add_paragraph()
    div_p.paragraph_format.space_before = Pt(2)
    div_p.paragraph_format.space_after  = Pt(4)
    pPr = div_p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'12')
    bot.set(qn('w:space'),'1'); bot.set(qn('w:color'),'0f172a')
    pBdr.append(bot); pPr.append(pBdr)

    # ── SUMMARY ──
    if cv_data.get("summary"):
        add_section_heading(h[1])
        add_para(cv_data["summary"].replace('\n',' ').strip(), size=10.5)

    # ── EDUCATION ──
    if cv_data.get("education"):
        add_section_heading(h[2])
        edu = re.sub(r'\s*•\s*', '\n• ', cv_data["education"])
        edu_blocks = re.split(r'\n(?=[A-Z\u0590-\u05FF\u0600-\u06FF])', edu)
        for block in edu_blocks:
            lines = [l.strip() for l in block.split('\n') if l.strip()]
            if not lines: continue
            raw_title = lines[0].lstrip('•-* ')
            title, date = parse_date_from_title(raw_title)
            entry_text = f"{title}   {date}" if date else title
            add_para(entry_text, bold=True, size=10.5, space_before=4, space_after=1)
            for l in lines[1:]:
                clean = l.lstrip('•-* ').strip()
                if clean: add_para(clean, italic=True, size=10, space_after=1)

    # ── SKILLS ──
    raw_skills = cv_data.get("skills","")
    if raw_skills:
        raw_skills_clean = re.sub(
            r'(Technical Skills|Business & Professional Skills|Professional Skills|'
            r'כישורים טכניים|כישורים מקצועיים|المهارات التقنية|المهارات المهنية)\s*[,\n]?',
            '', raw_skills, flags=re.IGNORECASE)
        skills_data = group_skills_segregated(raw_skills_clean)
        skills_title = h[8] if len(h) > 8 else ("כישורים" if output_lang=="Hebrew" else "المهارات" if output_lang=="Arabic" else "Skills")
        add_section_heading(skills_title)
        if skills_data["TECHNICAL"]:
            tech_label = "תכנות וכלים" if output_lang=="Hebrew" else "البرمجة والأدوات" if output_lang=="Arabic" else "Technical"
            p = doc.add_paragraph()
            set_para_format(p, is_rtl)
            r1 = p.add_run(f"{tech_label}: "); r1.bold=True; r1.font.size=Pt(10); r1.font.name=font_name
            r2 = p.add_run(" · ".join(skills_data["TECHNICAL"])); r2.font.size=Pt(10); r2.font.name=font_name
        if skills_data["PROFESSIONAL"]:
            prof_label = "כישורים מקצועיים" if output_lang=="Hebrew" else "المهارات المهنية" if output_lang=="Arabic" else "Professional"
            p = doc.add_paragraph()
            set_para_format(p, is_rtl)
            r1 = p.add_run(f"{prof_label}: "); r1.bold=True; r1.font.size=Pt(10); r1.font.name=font_name
            r2 = p.add_run(" · ".join(skills_data["PROFESSIONAL"])); r2.font.size=Pt(10); r2.font.name=font_name

    # ── PROJECTS ──
    if cv_data.get("projects"):
        add_section_heading(h[4])
        for block in split_cv_entries(cv_data["projects"]):
            block = re.sub(r'\s*•\s*', '\n• ', block)
            lines = [l.strip() for l in block.split('\n') if l.strip()]
            if not lines: continue
            raw_title = lines[0].lstrip('•-* ')
            title, date = parse_date_from_title(raw_title)
            entry_text = f"{title}   {date}" if date else title
            add_para(entry_text, bold=True, size=10.5, space_before=4, space_after=1)
            for line in lines[1:]:
                clean = line.lstrip('•-* ').strip()
                if not clean: continue
                if line.startswith(('•','-','*')): add_bullet_item(clean)
                else: add_para(clean, italic=True, size=10, space_after=1)

    # ── EXPERIENCE ──
    if cv_data.get("experience"):
        add_section_heading(h[3])
        for block in split_cv_entries(cv_data["experience"]):
            block = re.sub(r'\s*•\s*', '\n• ', block)
            lines = [l.strip() for l in block.split('\n') if l.strip()]
            if not lines: continue
            raw_title = lines[0].lstrip('•-* ')
            title, date = parse_date_from_title(raw_title)
            entry_text = f"{title}   {date}" if date else title
            add_para(entry_text, bold=True, size=10.5, space_before=4, space_after=1)
            sub_done = False
            for line in lines[1:]:
                clean = line.lstrip('•-* ').strip()
                if not clean: continue
                if line.startswith(('•','-','*')): add_bullet_item(clean)
                elif not sub_done:
                    add_para(clean, italic=True, size=10, space_after=1); sub_done=True
                else: add_bullet_item(clean)

    # ── VOLUNTEERING ──
    if cv_data.get("volunteering"):
        add_section_heading(h[6])
        vol = re.sub(r'^\s*[&\*]\s*$','', cv_data["volunteering"], flags=re.MULTILINE)
        for block in split_cv_entries(vol):
            block = re.sub(r'\s*•\s*', '\n• ', block)
            lines = [l.strip() for l in block.split('\n') if l.strip()]
            if not lines: continue
            raw_title = lines[0].lstrip('•-* ')
            title, date = parse_date_from_title(raw_title)
            add_para(f"{title}   {date}" if date else title, bold=True, size=10.5, space_before=4, space_after=1)
            for line in lines[1:]:
                clean = line.lstrip('•-* ').strip()
                if clean and line.startswith(('•','-','*')): add_bullet_item(clean)
                elif clean: add_para(clean, size=10, space_after=1)

    # ── LANGUAGES ──
    if cv_data.get("languages"):
        add_section_heading(h[7])
        langs_clean = re.sub(r'\s*•\s*', ' · ', cv_data["languages"]).strip(' ·')
        add_para(langs_clean, size=10)

    # ── COURSES ──
    if cv_data.get("courses_training"):
        add_section_heading(h[5])
        items = [c.strip() for c in re.split(r'[,\n]', cv_data["courses_training"]) if c.strip() and len(c.strip()) > 3]
        for item in items: add_bullet_item(item)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

    is_rtl = output_lang in ["Hebrew", "Arabic"]
    h = HEADINGS_MAP[output_lang]

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Arial' if is_rtl else 'Calibri'
    style.font.size = Pt(11)

    def set_rtl_para(para):
        """Apply RTL formatting to a paragraph."""
        if is_rtl:
            pPr = para._p.get_or_add_pPr()
            bidi = OxmlElement('w:bidi')
            pPr.append(bidi)
            jc = OxmlElement('w:jc')
            jc.set(qn('w:val'), 'right')
            pPr.append(jc)

    def add_para(text, bold=False, italic=False, size=11, align="left", color=None):
        p = doc.add_paragraph()
        if align == "center":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "right" or is_rtl:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_rtl_para(p)
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor(*color)
        return p

    def add_section_heading(title):
        p = add_para(title.upper(), bold=True, size=10, color=(30, 41, 59))
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(2)
        # Add bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '94a3b8')
        pBdr.append(bottom)
        pPr.append(pBdr)
        return p

    def add_bullet(text):
        p = doc.add_paragraph(style='List Bullet')
        set_rtl_para(p)
        if is_rtl:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(text.lstrip('•-* ').strip())
        run.font.size = Pt(10)
        return p

    # ── NAME & CONTACT ──
    name, contact = format_contact_info(cv_data.get("personal_details",""))
    if not name: name = "CANDIDATE"
    name_p = add_para(name.upper(), bold=True, size=20, align="center")
    name_p.paragraph_format.space_after = Pt(2)
    if contact:
        contact_clean = re.sub(r'(Email|Phone|Tel)\s*:\s*', '', contact, flags=re.IGNORECASE)
        add_para(contact_clean, size=10, align="center")

    # ── SECTIONS ──
    section_order = [
        ("summary",          h[1]),
        ("education",        h[2]),
        ("experience",       h[3]),
        ("projects",         h[4]),
        ("courses_training", h[5]),
        ("volunteering",     h[6]),
        ("languages",        h[7]),
    ]

    for key, heading in section_order:
        content = cv_data.get(key, "").strip()
        if not content:
            continue

        add_section_heading(heading)

        if key in ["experience", "projects", "volunteering"]:
            for block in split_cv_entries(content):
                block = re.sub(r'\s*•\s*', '\n• ', block)
                lines = [l.strip() for l in block.split('\n') if l.strip()]
                for i, line in enumerate(lines):
                    clean = line.lstrip('•-* ').strip()
                    if not clean:
                        continue
                    if i == 0:
                        add_para(clean, bold=True, size=10.5)
                    elif line.startswith(('•', '-', '*')):
                        add_bullet(clean)
                    else:
                        add_para(clean, italic=True, size=10)

        elif key == "education":
            # Each school on its own
            edu_entries = re.split(r'\n(?=[A-Z])', re.sub(r'\s*•\s*', '\n• ', content))
            for entry in edu_entries:
                lines = [l.strip() for l in entry.split('\n') if l.strip()]
                if not lines:
                    continue
                add_para(lines[0].lstrip('•-* '), bold=True, size=10.5)
                for l in lines[1:]:
                    add_para(l.lstrip('•-* '), italic=True, size=10)

        elif key == "courses_training":
            items = [c.strip() for c in re.split(r'[,\n]', content) if c.strip() and len(c.strip()) > 3]
            for item in items:
                add_bullet(item)

        elif key == "languages":
            langs_clean = re.sub(r'\s*•\s*', ' · ', content).strip(' ·')
            add_para(langs_clean, size=10)

        elif key == "summary":
            summary = content.replace('\n', ' ').strip()
            add_para(summary, size=10.5)

        else:
            add_para(content, size=10)

    # ── SKILLS ──
    raw_skills = cv_data.get("skills", "")
    if raw_skills:
        raw_skills = re.sub(r'(Technical Skills|Business & Professional Skills|Professional Skills)\s*[,\n]?', '', raw_skills, flags=re.IGNORECASE)
        skills_data = group_skills_segregated(raw_skills)

        if skills_data["TECHNICAL"]:
            add_section_heading(h[8] if len(h) > 8 else "Technical Skills")
            add_para(" · ".join(skills_data["TECHNICAL"]), size=10)
        if skills_data["PROFESSIONAL"]:
            add_section_heading("Professional Skills")
            add_para(" · ".join(skills_data["PROFESSIONAL"]), size=10)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ============================================================
# NAVIGATION HELPERS
# ============================================================

TAB_NAMES = ["The Job", "CV Content", "Analysis", "Refinement", "Preview", "Export"]
TAB_ICONS = ["🎯", "📄", "📊", "✏️", "👁️", "⬇️"]

def go_to(tab_index):
    st.session_state.active_tab = tab_index
    if tab_index < 2:
        # Back to Job or CV Content — wipe deep analysis so it re-runs
        st.session_state.ai_cv_analysis = {}
        st.session_state.analysis_done = False
    if tab_index < 3:
        # Back before Refinement — wipe dynamic questions
        st.session_state.dynamic_questions = []
    if tab_index < 4:
        # Back before Preview — wipe rewrite output
        st.session_state.cv_adjusted = False
        st.session_state.adjusted_cv_data = {}
        st.session_state.final_cv_data = {}
        st.session_state.improvement_log = []
    st.rerun()

LOGO_SVG = """<svg width="32" height="32" viewBox="0 0 32 32" fill="none" xmlns="http://www.w3.org/2000/svg">
  <rect x="5" y="2" width="18" height="24" rx="2" stroke="#0F172A" stroke-width="1.8" fill="white"/>
  <line x1="9" y1="9" x2="19" y2="9" stroke="#E5E7EB" stroke-width="1.2"/>
  <line x1="9" y1="13" x2="19" y2="13" stroke="#E5E7EB" stroke-width="1.2"/>
  <line x1="9" y1="17" x2="15" y2="17" stroke="#E5E7EB" stroke-width="1.2"/>
  <circle cx="22" cy="22" r="7" fill="white" stroke="#0F172A" stroke-width="1.8"/>
  <line x1="22" y1="17" x2="22" y2="22" stroke="#0F172A" stroke-width="2" stroke-linecap="round"/>
  <line x1="22" y1="22" x2="25.5" y2="25.5" stroke="#2563EB" stroke-width="2" stroke-linecap="round"/>
  <circle cx="22" cy="22" r="1.2" fill="#0F172A"/>
</svg>"""

LOGO_SVG_SM = """<svg width="24" height="24" viewBox="0 0 32 32" fill="none" xmlns="http://www.w3.org/2000/svg">
  <rect x="5" y="2" width="18" height="24" rx="2" stroke="#0F172A" stroke-width="1.8" fill="white"/>
  <circle cx="22" cy="22" r="7" fill="white" stroke="#0F172A" stroke-width="1.8"/>
  <line x1="22" y1="17" x2="22" y2="22" stroke="#0F172A" stroke-width="2" stroke-linecap="round"/>
  <line x1="22" y1="22" x2="25.5" y2="25.5" stroke="#2563EB" stroke-width="2" stroke-linecap="round"/>
  <circle cx="22" cy="22" r="1.2" fill="#0F172A"/>
</svg>"""


def render_header():
    st.markdown(f"""
    <div class="rp-header">
        {LOGO_SVG}
        <span class="rp-brand">ResumePilot</span>
        <span class="rp-tagline">Land more interviews.</span>
    </div>
    """, unsafe_allow_html=True)


def render_step_bar(current):
    html = '<div class="rp-step-bar">'
    for i, name in enumerate(TAB_NAMES):
        if i < current:
            dot_cls, name_cls, icon = "rp-dot-done", "rp-name-done", "✓"
        elif i == current:
            dot_cls, name_cls, icon = "rp-dot-active", "rp-name-active", str(i+1)
        else:
            dot_cls, name_cls, icon = "rp-dot-todo", "rp-name-todo", str(i+1)
        html += f'<div class="rp-step-item"><div class="rp-step-dot {dot_cls}">{icon}</div><span class="rp-step-name {name_cls}">{name}</span></div>'
        if i < len(TAB_NAMES) - 1:
            conn = "rp-conn-done" if i < current else "rp-conn-todo"
            html += f'<div class="rp-connector {conn}"></div>'
    html += '</div>'
    st.markdown(html, unsafe_allow_html=True)


def render_landing():
    """Full ResumePilot landing page shown before user starts."""

    # HERO
    st.markdown("""
<div style="background:radial-gradient(ellipse 80% 45% at 50% -10%, #EFF6FF 0%, #F8FAFC 65%);padding-bottom:8px;">
  <div class="rp-hero">
    <h1>Turn any CV into a<br><span>job-targeted application.</span></h1>
    <p>ResumePilot analyzes the job description, finds what your CV is missing, asks targeted questions, and creates a stronger resume without inventing experience.</p>
  </div>
</div>
    """, unsafe_allow_html=True)

    # CTA BUTTON — real Streamlit widget
    st.markdown("<div style='height:24px;'></div>", unsafe_allow_html=True)
    col1, col2, col3 = st.columns([2, 1.2, 2])
    with col2:
        if st.button("Start tailoring →", type="primary", use_container_width=True):
            st.session_state.show_landing = False
            st.rerun()
    st.markdown("<div style='height:16px;'></div>", unsafe_allow_html=True)

    # BENTO GRID — 6 features
    st.markdown("""
<div class="rp-bento">
  <div class="rp-bento-cell">
    <div class="rp-bento-icon">🔍</div>
    <div class="rp-bento-title">Semantic job matching</div>
    <div class="rp-bento-desc">We don't just scan for keywords — we understand what the role actually requires and where your CV falls short.</div>
  </div>
  <div class="rp-bento-cell">
    <div class="rp-bento-icon">🎯</div>
    <div class="rp-bento-title">Gap detection</div>
    <div class="rp-bento-desc">Every missing ATS keyword, missing metric, and weak bullet is flagged clearly so you know exactly what to fix.</div>
  </div>
  <div class="rp-bento-cell">
    <div class="rp-bento-icon">💬</div>
    <div class="rp-bento-title">Targeted refinement questions</div>
    <div class="rp-bento-desc">AI generates 5–8 personalized questions based on your actual CV gaps — not generic prompts.</div>
  </div>
  <div class="rp-bento-cell">
    <div class="rp-bento-icon">✍️</div>
    <div class="rp-bento-title">Honest AI rewriting</div>
    <div class="rp-bento-desc">Your CV is rewritten using only what you told us — your real experience, your actual projects, your words.</div>
  </div>
  <div class="rp-bento-cell">
    <div class="rp-bento-icon">🌐</div>
    <div class="rp-bento-title">English / Hebrew / Arabic output</div>
    <div class="rp-bento-desc">Professional localization via Claude — not word-by-word translation. Native phrasing for the Israeli job market.</div>
  </div>
  <div class="rp-bento-cell">
    <div class="rp-bento-icon">📁</div>
    <div class="rp-bento-title">DOCX export</div>
    <div class="rp-bento-desc">Download a formatted, recruiter-ready DOCX or TXT file — ready to attach and send immediately.</div>
  </div>
</div>
    """, unsafe_allow_html=True)

    # HOW IT WORKS
    st.markdown("""
<div class="rp-how">
  <div class="rp-how-inner">
    <div class="rp-how-label">How it works</div>
    <div class="rp-how-title">From job description to tailored CV in minutes.</div>
    <div class="rp-how-steps">
      <div class="rp-how-step">
        <div class="rp-how-num">1</div>
        <div class="rp-how-step-title">Paste the job</div>
        <div class="rp-how-step-desc">Drop in any job description from LinkedIn, a company site, or anywhere else.</div>
      </div>
      <div class="rp-how-step">
        <div class="rp-how-num">2</div>
        <div class="rp-how-step-title">Upload your CV</div>
        <div class="rp-how-step-desc">We analyze it, score the match, and pinpoint every gap worth addressing.</div>
      </div>
      <div class="rp-how-step">
        <div class="rp-how-num">3</div>
        <div class="rp-how-step-title">Download &amp; apply</div>
        <div class="rp-how-step-desc">Get a tailored resume in English, Hebrew, or Arabic — ready to send.</div>
      </div>
    </div>
  </div>
</div>
    """, unsafe_allow_html=True)

    # FOOTER
    st.markdown("""
<div class="rp-footer">
  <div class="rp-footer-brand">ResumePilot</div>
  <p class="rp-footer-p">AI-powered CV tailoring for the modern job market.</p>
</div>
    """, unsafe_allow_html=True)


def nav_buttons(current, can_proceed=True, proceed_label="Continue →", back_label="← Back"):
    st.markdown("<div style='margin-top:28px; padding-top:8px; border-top:1px solid #E5E7EB;'></div>", unsafe_allow_html=True)
    col1, col2, col3 = st.columns([1, 6, 1])
    with col1:
        if current > 0:
            if st.button(back_label, key=f"back_{current}", use_container_width=True):
                go_to(current - 1)
    with col3:
        if current < len(TAB_NAMES) - 1:
            btn = st.button(
                proceed_label,
                key=f"next_{current}",
                use_container_width=True,
                disabled=not can_proceed,
                type="primary"
            )
            if btn:
                go_to(current + 1)


# ============================================================
# APP HEADER + LANDING PAGE
# ============================================================

# Show landing page if not started
if st.session_state.show_landing:
    render_landing()
    st.stop()

# Otherwise show the app
render_header()
current_tab = st.session_state.active_tab
render_step_bar(current_tab)
st.markdown("<hr style='border:none;border-top:1px solid #E5E7EB;margin-bottom:28px;'>", unsafe_allow_html=True)


# ============================================================
# TAB 0 — THE JOB
# ============================================================

if current_tab == 0:
    st.markdown('<div class="rp-page-badge">STEP 1 OF 6</div>', unsafe_allow_html=True)
    st.markdown('<div class="rp-page-title">Target Position</div>', unsafe_allow_html=True)
    st.markdown('<div class="rp-page-sub">Tell us the role you\'re going for. We\'ll use it to match and tailor everything.</div>', unsafe_allow_html=True)

    st.session_state.job_role = st.text_input(
        "Target Job Title",
        value=st.session_state.job_role,
        placeholder="e.g., Data Analyst Intern, Software Engineer, Marketing Associate"
    )

    st.session_state.job_desc = st.text_area(
        "Job Description",
        value=st.session_state.job_desc,
        height=280,
        placeholder="Paste the full job description here — include responsibilities, requirements, and qualifications..."
    )

    col_demo, _ = st.columns([1, 3])
    with col_demo:
        if st.button("Load demo job", use_container_width=True):
            st.session_state.job_role = "Junior Data Analyst"
            st.session_state.job_desc = """We are looking for a Junior Data Analyst to join our growing team.

Responsibilities:
- Analyze large datasets using Python and SQL
- Build dashboards and reports using Tableau or Power BI
- Apply statistical methods including regression and modeling
- Collaborate with cross-functional teams to deliver insights
- Present findings to stakeholders

Requirements:
- Bachelor's degree in Statistics, Mathematics, Computer Science or related field
- Proficiency in Python and SQL
- Experience with data visualization tools (Tableau, Power BI, or Matplotlib)
- Strong analytical and problem-solving skills
- Knowledge of machine learning concepts
- Git version control experience
- Excellent communication and teamwork skills
- Hebrew and English required"""
            st.rerun()

    # Validation feedback
    job_ready = bool(st.session_state.job_role.strip() and st.session_state.job_desc.strip())
    if st.session_state.job_desc.strip():
        is_valid, err = validate_job_description(st.session_state.job_desc)
        if not is_valid:
            st.error(f"⚠️ {err}")
            job_ready = False
        else:
            st.markdown('<div class="alert-box alert-green">✓ Job description looks good.</div>', unsafe_allow_html=True)

    nav_buttons(0, can_proceed=job_ready, proceed_label="Continue →")


# ============================================================
# TAB 1 — CV CONTENT
# ============================================================

elif current_tab == 1:
    st.markdown('<div class="rp-page-badge">STEP 2 OF 6</div>', unsafe_allow_html=True)
    st.markdown('<div class="rp-page-title">CV Content</div>', unsafe_allow_html=True)
    st.markdown('<div class="rp-page-sub">Upload or paste your resume in English. Hebrew and Arabic output available in Preview.</div>', unsafe_allow_html=True)

    # English-only notice
    st.info("Please enter your CV in English. You can choose Hebrew or Arabic output in the Preview tab.")

    method = st.radio(
        "How would you like to provide your CV?",
        ["Upload File (PDF/DOCX/TXT)", "Paste Text", "Manual Entry"],
        horizontal=True,
        key="cv_method_radio"
    )
    st.session_state.input_method = method

    st.markdown("---")

    # ---- UPLOAD ----
    if method == "Upload File (PDF/DOCX/TXT)":
        uploaded_file = st.file_uploader("Upload your CV", type=["pdf","docx","txt"], label_visibility="collapsed")
        if uploaded_file:
            with st.spinner("Reading your CV..."):
                if uploaded_file.type == "application/pdf":
                    text = extract_text_from_pdf(uploaded_file)
                elif "wordprocessingml" in uploaded_file.type:
                    text = extract_text_from_docx(uploaded_file)
                else:
                    text = str(uploaded_file.read(), "utf-8")

            if text.strip():
                # Detect if uploaded CV is in Hebrew or Arabic
                detected_lang = detect_language(text)
                if detected_lang in ["Hebrew", "Arabic"]:
                    st.warning(f"⚠️ Your uploaded CV appears to be in **{detected_lang}**. For best results, please upload an English version of your CV. The app needs English text for accurate analysis and job matching.")
                    st.markdown("**You can still continue**, but the analysis accuracy may be lower.")

                st.session_state.cv_full_text = text
                st.session_state.manual_cv_data = classify_sections_refined(text)
                st.session_state.cv_uploaded = True
                st.success(f"✅ CV uploaded and parsed! ({len(text.split())} words detected)")

                with st.expander("👀 Preview detected sections"):
                    for k, v in st.session_state.manual_cv_data.items():
                        if v.strip():
                            st.markdown(f"**{k.replace('_',' ').title()}**")
                            st.text(v[:300] + ("..." if len(v) > 300 else ""))
            else:
                st.error("❌ Could not extract text from this file. Try copy-pasting instead.")

    # ---- PASTE TEXT ----
    elif method == "Paste Text":
        pasted = st.text_area(
            "Paste your CV text here (in English):",
            value=st.session_state.cv_full_text,
            height=320,
            placeholder="Copy and paste your full CV content here in English..."
        )
        st.session_state.cv_full_text = pasted

        if pasted.strip():
            detected_lang = detect_language(pasted)
            if detected_lang in ["Hebrew", "Arabic"]:
                st.warning(f"⚠️ This text appears to be in **{detected_lang}**. Please paste your CV in English for best results.")

        if st.button("🔍 Auto-Structure Text", type="primary"):
            if pasted.strip():
                st.session_state.manual_cv_data = classify_sections_refined(pasted)
                st.session_state.cv_uploaded = True
                st.success("✅ Text structured into sections!")
                with st.expander("👀 Preview detected sections"):
                    for k, v in st.session_state.manual_cv_data.items():
                        if v.strip():
                            st.markdown(f"**{k.replace('_',' ').title()}**")
                            st.text(v[:300] + ("..." if len(v) > 300 else ""))
            else:
                st.warning("Please paste your CV text first.")

        if st.session_state.cv_full_text.strip():
            st.session_state.cv_uploaded = True

    # ---- MANUAL ENTRY ----
    elif method == "Manual Entry":
        st.markdown("**Fill in each section in English below:**")
        fields = [
            ("personal_details", "👤 Full Name & Contact Info",
             "Your Full Name\nyour.email@example.com | +000-000-0000 | City, Country | linkedin.com/in/yourprofile"),
            ("summary",          "💡 Professional Summary",
             "e.g., Final-year Computer Science student with experience in data analysis and software development. Passionate about turning data into actionable insights."),
            ("education",        "🎓 Education",
             "e.g., University Name\nBSc in Your Major\n2022–2025 | GPA: 3.8"),
            ("experience",       "💼 Work Experience",
             "e.g., Job Title – Company Name\nMonth Year – Month Year\n- What you did and achieved\n- Another responsibility or result"),
            ("projects",         "🚀 Projects",
             "e.g., Project Name – Brief description\n- Tools used: Python, SQL, etc.\n- What problem it solved or result it produced"),
            ("skills",           "🛠️ Skills",
             "e.g., Python, SQL, Excel, Git, Data Analysis, Communication, Problem Solving, Teamwork"),
            ("courses_training", "📚 Courses & Certifications",
             "e.g., Google Data Analytics Certificate (Coursera, 2024)\nAdvanced Excel for Data Analysis (LinkedIn Learning)"),
            ("languages",        "🌐 Languages",
             "e.g., English (Native), Arabic (Fluent), French (Intermediate)"),
            ("volunteering",     "🤝 Volunteering & Activities",
             "e.g., Volunteer Tutor – Local NGO | 2023–Present\n- Helped 20+ students with math and science\nUniversity Student Council – Events Coordinator"),
        ]
        c1, c2 = st.columns(2)
        any_filled = False
        for i, (key, label, placeholder) in enumerate(fields):
            col = c1 if i % 2 == 0 else c2
            with col:
                st.markdown(f"**{label}**")
                val = st.text_area(
                    label, value=st.session_state.manual_cv_data.get(key,""),
                    placeholder=placeholder, key=f"manual_{key}",
                    label_visibility="collapsed", height=110
                )
                st.session_state.manual_cv_data[key] = val
                if val.strip(): any_filled = True

        if any_filled:
            st.session_state.cv_full_text = "\n".join(st.session_state.manual_cv_data.values())
            st.session_state.cv_uploaded = True

    if method == "Paste Text":
        pasted_text = st.session_state.cv_full_text.strip()
        if not pasted_text:
            st.error("⚠️ Please paste your CV text before continuing.")
            cv_ready = False
        elif len(pasted_text) < 50:
            st.error(f"⚠️ CV text is too short ({len(pasted_text)} characters). Please paste your full resume — minimum 50 characters.")
            cv_ready = False
        else:
            cv_ready = True
    else:
        cv_ready = bool(
            st.session_state.cv_full_text.strip() or
            any(v.strip() for v in st.session_state.manual_cv_data.values())
        )
    nav_buttons(1, can_proceed=cv_ready, proceed_label="Continue →")


# ============================================================
# TAB 2 — ANALYSIS
# ============================================================

elif current_tab == 2:
    st.markdown('<div class="rp-page-badge">STEP 3 OF 6</div>', unsafe_allow_html=True)
    st.markdown('<div class="rp-page-title">Match Analysis</div>', unsafe_allow_html=True)
    st.markdown('<div class="rp-page-sub">Deep AI analysis of your CV against this role — strengths, gaps, ATS keywords, and quick wins.</div>', unsafe_allow_html=True)

    cv_text = st.session_state.cv_full_text or "\n".join(st.session_state.manual_cv_data.values())
    job_ok  = bool(st.session_state.job_desc.strip())
    cv_ok   = bool(cv_text.strip())

    if not job_ok or not cv_ok:
        st.warning("Please complete The Job and CV Content steps first.")
        nav_buttons(2, can_proceed=False)
    else:
        is_valid, err = validate_job_description(st.session_state.job_desc)
        if not is_valid:
            st.error(f"⚠️ {err}")
            nav_buttons(2, can_proceed=False)
        else:
            # ── Parse CV into sections ──
            cv_sections = st.session_state.manual_cv_data.copy()
            if not any(v.strip() for v in cv_sections.values()):
                cv_sections = classify_sections_refined(cv_text)

            # ── Run deep Claude analysis (Step 1+2) ──
            if not st.session_state.ai_cv_analysis:
                with st.spinner("🤖 Analyzing your CV against the job requirements... (~15 seconds)"):
                    result = call_claude_deep_analysis(
                        cv_sections,
                        st.session_state.job_role,
                        st.session_state.job_desc,
                    )
                if result:
                    st.session_state.ai_cv_analysis = result
                else:
                    # Fallback: run rule-based and build a compatible structure
                    rule = run_rule_based_analysis(cv_text, st.session_state.job_desc)
                    st.session_state.ai_cv_analysis = {
                        "cv_strengths":           [g["label"] for g in rule["strong"]],
                        "cv_weaknesses":          [g["label"] for g in rule["missing"]],
                        "hard_skills_required":   [g["label"] for g in rule["missing"]],
                        "soft_skills_required":   [],
                        "domain_knowledge_required": [],
                        "ats_keywords_missing":   [g["label"] for g in rule["missing"]],
                        "ats_keywords_present":   [g["label"] for g in rule["strong"]],
                        "match_score":            rule["score"],
                        "match_label":            ("Strong Match" if rule["score"] >= 70
                                                   else "Moderate Match" if rule["score"] >= 45
                                                   else "Weak Match"),
                        "score_rationale":        "Based on rule-based keyword matching (AI unavailable).",
                        "metrics_missing":        [],
                        "quick_wins":             [f"Add {g['label']} to your CV" for g in rule["missing"][:3]],
                    }
                    if get_anthropic_api_key() is None:
                        st.error("⚠️ **ANTHROPIC_API_KEY not found.** Add it to `.streamlit/secrets.toml` to enable AI analysis. Showing rule-based results.")
                    else:
                        st.warning("⚠️ AI analysis failed (API error). Showing rule-based results.")

            # Persist to analysis_results for downstream tabs
            ai = st.session_state.ai_cv_analysis
            st.session_state.analysis_results = {
                "matches": ai.get("ats_keywords_present", []),
                "missing": ai.get("ats_keywords_missing", []),
                "score":   ai.get("match_score", 0),
            }
            st.session_state.analysis_done = True

            score = ai.get("match_score", 0)
            score_color = "#16a34a" if score >= 70 else "#ea580c" if score >= 45 else "#dc2626"

            # ── Score Card ──
            st.markdown(f"""
            <div class="content-card" style="text-align:center; padding:2rem;">
                <div style="font-size:5rem;font-weight:900;color:{score_color};line-height:1;">{score}%</div>
                <div style="font-size:1.3rem;font-weight:700;color:{score_color};margin:8px 0;">{ai.get("match_label","")}</div>
                <div style="color:#64748b;font-size:0.95rem;max-width:520px;margin:0 auto;line-height:1.6;">
                    {ai.get("score_rationale","")}
                </div>
            </div>
            """, unsafe_allow_html=True)

            st.markdown("---")

            # ── Strengths and Weaknesses ──
            col_s, col_w = st.columns(2)
            strengths  = ai.get("cv_strengths", [])
            weaknesses = ai.get("cv_weaknesses", [])

            with col_s:
                st.markdown("""<div style="background:#f0fdf4;border:1px solid #bbf7d0;
                    border-radius:12px;padding:1.2rem;">
                    <div style="font-weight:700;color:#16a34a;font-size:1rem;margin-bottom:0.8rem;">
                    ✅ CV Strengths</div>""", unsafe_allow_html=True)
                if strengths:
                    for s in strengths:
                        st.markdown(f'<div style="margin-bottom:6px;font-size:0.88rem;color:#166534;">✓ {s}</div>', unsafe_allow_html=True)
                else:
                    st.markdown('<div style="color:#6b7280;font-size:0.85rem;">No clear strengths identified for this role.</div>', unsafe_allow_html=True)
                st.markdown('</div>', unsafe_allow_html=True)

            with col_w:
                st.markdown("""<div style="background:#fff1f2;border:1px solid #fecdd3;
                    border-radius:12px;padding:1.2rem;">
                    <div style="font-weight:700;color:#be123c;font-size:1rem;margin-bottom:0.8rem;">
                    ⚠️ Gaps to Address</div>""", unsafe_allow_html=True)
                if weaknesses:
                    for w in weaknesses:
                        st.markdown(f'<div style="margin-bottom:6px;font-size:0.88rem;color:#9f1239;">✗ {w}</div>', unsafe_allow_html=True)
                else:
                    st.markdown('<div style="color:#16a34a;font-size:0.85rem;font-weight:600;">🎉 No major gaps!</div>', unsafe_allow_html=True)
                st.markdown('</div>', unsafe_allow_html=True)

            st.markdown("---")

            # ── ATS Keywords ──
            ats_missing  = ai.get("ats_keywords_missing", [])
            ats_present  = ai.get("ats_keywords_present", [])
            metrics_miss = ai.get("metrics_missing", [])

            st.markdown("**🔍 ATS Keyword Analysis**")
            col_ap, col_am = st.columns(2)
            with col_ap:
                if ats_present:
                    st.markdown('<div style="background:#f0fdf4;border-radius:10px;padding:1rem;border:1px solid #bbf7d0;">', unsafe_allow_html=True)
                    st.markdown('<div style="font-weight:700;color:#16a34a;font-size:0.9rem;margin-bottom:0.5rem;">✅ Keywords Present</div>', unsafe_allow_html=True)
                    chips = " ".join(
                        f'<span style="display:inline-block;background:#dcfce7;color:#166534;padding:3px 10px;border-radius:9999px;font-size:0.8rem;margin:2px;">{k}</span>'
                        for k in ats_present[:12]
                    )
                    st.markdown(chips, unsafe_allow_html=True)
                    st.markdown('</div>', unsafe_allow_html=True)

            with col_am:
                if ats_missing:
                    st.markdown('<div style="background:#fff1f2;border-radius:10px;padding:1rem;border:1px solid #fecdd3;">', unsafe_allow_html=True)
                    st.markdown('<div style="font-weight:700;color:#be123c;font-size:0.9rem;margin-bottom:0.5rem;">❌ Keywords Missing</div>', unsafe_allow_html=True)
                    chips = " ".join(
                        f'<span style="display:inline-block;background:#ffe4e6;color:#9f1239;padding:3px 10px;border-radius:9999px;font-size:0.8rem;margin:2px;">{k}</span>'
                        for k in ats_missing[:12]
                    )
                    st.markdown(chips, unsafe_allow_html=True)
                    st.markdown('</div>', unsafe_allow_html=True)

            # ── Bullets needing metrics ──
            if metrics_miss:
                st.markdown("---")
                st.markdown("**📊 Bullets That Need Metrics (add numbers = more impact)**")
                for m in metrics_miss:
                    st.markdown(f'<div style="background:#fffbeb;border-left:3px solid #f59e0b;padding:8px 12px;border-radius:0 6px 6px 0;margin-bottom:6px;font-size:0.88rem;color:#78350f;">{m}</div>', unsafe_allow_html=True)

            # ── Quick Wins ──
            quick_wins = ai.get("quick_wins", [])
            if quick_wins:
                st.markdown("---")
                st.markdown("**⚡ Quick Wins — biggest impact improvements**")
                for i, qw in enumerate(quick_wins, 1):
                    st.markdown(f'<div style="background:#f0f9ff;border-left:3px solid #0ea5e9;padding:8px 12px;border-radius:0 6px 6px 0;margin-bottom:6px;font-size:0.88rem;color:#0c4a6e;">{i}. {qw}</div>', unsafe_allow_html=True)

            st.markdown("---")
            nav_buttons(2, can_proceed=True, proceed_label="Next: Refine CV →")


# ============================================================
# TAB 3 — REFINEMENT
# ============================================================

elif current_tab == 3:
    st.markdown('<div class="rp-page-badge">STEP 4 OF 6</div>', unsafe_allow_html=True)
    st.markdown('<div class="rp-page-title">Refine Your CV</div>', unsafe_allow_html=True)
    st.markdown('<div class="rp-page-sub">Answer these personalized questions to unlock a stronger, more specific CV rewrite.</div>', unsafe_allow_html=True)

    if not st.session_state.analysis_done:
        st.info("ℹ️ Please complete the Analysis step first.")
        nav_buttons(3, can_proceed=False)
    else:
        score = st.session_state.analysis_results.get("score", 0)
        if score >= 70:
            st.success(f"🎉 Your CV scores **{score}%** — strong match. Answer below to make it even sharper.")
        elif score >= 40:
            st.warning(f"Your CV scores **{score}%**. These answers will help fill the gaps.")
        else:
            st.error(f"Your CV scores **{score}%**. Be specific in your answers — they'll significantly improve the rewrite.")

        st.markdown("---")

        # ── Generate dynamic personalized questions (Step 3) ──
        if not st.session_state.dynamic_questions:
            cv_sections = st.session_state.manual_cv_data.copy()
            if not any(v.strip() for v in cv_sections.values()):
                cv_sections = classify_sections_refined(
                    st.session_state.cv_full_text or ""
                )
            with st.spinner("🤖 Generating personalized questions for your CV... (~10 seconds)"):
                questions = call_claude_generate_questions(
                    cv_sections,
                    st.session_state.job_role,
                    st.session_state.job_desc,
                    st.session_state.ai_cv_analysis,
                )
            if questions:
                st.session_state.dynamic_questions = questions
            else:
                # Fallback: generic gap-based questions
                missing = st.session_state.analysis_results.get("missing", [])
                fallback_qs = []
                for i, gap in enumerate(missing[:5]):
                    gap_label = gap if isinstance(gap, str) else str(gap)
                    fallback_qs.append({
                        "key": f"q_gap_{i}",
                        "question": f"The role requires **{gap_label}**. Do you have any experience, coursework, or projects involving this? Describe briefly.",
                        "hint": f"e.g., I used {gap_label} in a university project / I studied it in [course] / I don't have this yet",
                        "section": "skills",
                        "gap_addressed": f"Clarifies exposure to {gap_label}",
                    })
                if not fallback_qs:
                    fallback_qs = [
                        {"key": "q_achievement", "question": "What is your biggest achievement relevant to this role? Include numbers or outcomes if possible.", "hint": "e.g., Built a model that improved accuracy by 15% / Led a team of 5 students for 6 months", "section": "experience", "gap_addressed": "Adds quantifiable achievement to rewrite"},
                        {"key": "q_tools", "question": "Are there tools, platforms, or technologies you use regularly that aren't listed on your CV?", "hint": "e.g., I use Tableau weekly but forgot to add it / I've been learning AWS on the side", "section": "skills", "gap_addressed": "Surfaces unlisted relevant skills"},
                    ]
                st.session_state.dynamic_questions = fallback_qs
                if get_anthropic_api_key() is None:
                    st.error("⚠️ **ANTHROPIC_API_KEY not found.** Showing fallback questions. Add your key to enable personalized AI questions.")
                else:
                    st.warning("⚠️ Could not generate personalized questions (API error). Showing standard questions.")

        # ── Render questions (Step 4: store all answers in structured memory) ──
        questions = st.session_state.dynamic_questions
        st.markdown(f"**Answer these {len(questions)} questions** — your answers are saved automatically and used to rewrite your CV:")
        st.markdown("")

        for q in questions:
            key    = q.get("key", "q_unknown")
            text   = q.get("question", "")
            hint   = q.get("hint", "")
            section= q.get("section", "")
            gap    = q.get("gap_addressed", "")

            # Show section badge if available
            badge = f'<span style="font-size:0.72rem;background:#e0e7ff;color:#3730a3;padding:2px 8px;border-radius:9999px;font-weight:600;margin-left:6px;text-transform:uppercase;">{section}</span>' if section else ""
            st.markdown(f'<div style="font-weight:700;font-size:0.95rem;color:#1e293b;margin-bottom:4px;">{text}{badge}</div>', unsafe_allow_html=True)
            if gap:
                st.markdown(f'<div style="font-size:0.78rem;color:#64748b;margin-bottom:6px;">Why we ask: {gap}</div>', unsafe_allow_html=True)

            # Retrieve prior answer (supports both old string format and new dict format)
            prior = st.session_state.follow_up_answers.get(key, "")
            if isinstance(prior, dict):
                prior = prior.get("answer", "")

            answer = st.text_area(
                text,
                value=prior,
                placeholder=hint,
                key=f"refine_{key}",
                label_visibility="collapsed",
                height=90,
            )

            # Store as structured memory (Step 4)
            st.session_state.follow_up_answers[key] = {
                "question":      text,
                "answer":        answer,
                "section":       section,
                "gap_addressed": gap,
            }

            st.markdown("")

        st.markdown("---")
        st.markdown('<div style="background:#f0f9ff;border:1px solid #bae6fd;border-radius:10px;padding:1rem;font-size:0.88rem;color:#0369a1;">'
                    '💡 <strong>Tip:</strong> Skip any question that doesn\'t apply — write "N/A" or leave it blank. '
                    'The AI rewriter will only use the answers you provide. It will never invent information you didn\'t give.'
                    '</div>', unsafe_allow_html=True)

        st.markdown("")
        nav_buttons(3, can_proceed=True, proceed_label="Next: Preview CV →")


# ============================================================
# TAB 4 — PREVIEW
# ============================================================

elif current_tab == 4:
    st.markdown('<div class="rp-page-badge">STEP 5 OF 6</div>', unsafe_allow_html=True)
    st.markdown('<div class="rp-page-title">Preview</div>', unsafe_allow_html=True)
    st.markdown('<div class="rp-page-sub">Your professionally rewritten, ATS-optimized resume. Review before downloading.</div>', unsafe_allow_html=True)

    lang = st.radio("Output language:", ["English","Hebrew","Arabic"], horizontal=True, key="lang_select")
    if lang != st.session_state.active_lang:
        st.session_state.active_lang = lang
        st.markdown(
            "<script>window.parent.document.querySelector('.main').scrollTo(0,0);</script>",
            unsafe_allow_html=True,
        )
        st.rerun()
    st.session_state.active_lang = lang

    # ── Get base CV sections ──
    base_data = st.session_state.manual_cv_data.copy()
    if not any(v.strip() for v in base_data.values()):
        base_data = classify_sections_refined(st.session_state.cv_full_text or "")

    has_content = any(v.strip() for v in base_data.values())

    if not has_content:
        st.warning("⚠️ No CV content found. Please go back to **CV Content** and add your information.")
        nav_buttons(4, can_proceed=False)
    else:
        # ── Steps 5-10: Full CV Rewrite (English, run once per session) ──
        if not st.session_state.cv_adjusted:
            with st.spinner("✨ Rewriting your CV for this role... (~20 seconds)"):
                rewrite_result = call_claude_rewrite_cv(
                    base_data,
                    st.session_state.job_role,
                    st.session_state.job_desc,
                    st.session_state.ai_cv_analysis,
                    st.session_state.follow_up_answers,
                )

            if rewrite_result:
                # Extract improvement_log from the rewrite response (Step 10)
                improvement_log = rewrite_result.pop("improvement_log", [])
                st.session_state.improvement_log = improvement_log if isinstance(improvement_log, list) else []

                st.session_state.adjusted_cv_data = build_final_cv(base_data, rewrite_result, st.session_state.job_role)
                st.session_state.cv_adjusted = True
            else:
                # Fallback: use original CV with minimal improvement
                fallback = base_data.copy()
                if st.session_state.job_role and fallback.get("summary"):
                    if st.session_state.job_role.lower() not in fallback["summary"].lower():
                        fallback["summary"] = (
                            f"Motivated professional targeting a {st.session_state.job_role} role. "
                            + fallback["summary"]
                        )
                st.session_state.adjusted_cv_data = fallback
                st.session_state.cv_adjusted = True
                st.session_state.improvement_log = []
                if get_anthropic_api_key() is None:
                    st.error(
                        "⚠️ **ANTHROPIC_API_KEY not found.** "
                        "Add it to `.streamlit/secrets.toml` as `ANTHROPIC_API_KEY = \"sk-ant-...\"` "
                        "to enable AI rewriting. Showing your original CV for now."
                    )
                else:
                    st.warning("⚠️ AI rewrite failed (API error). Showing your original CV.")

        english_cv = st.session_state.adjusted_cv_data or base_data

        # ── Step 9: Localization for Hebrew / Arabic ──
        if lang in ["Hebrew", "Arabic"]:
            cache_key = f"translated_{lang}"
            cached = st.session_state.get(cache_key, {})
            if not cached:
                with st.spinner(f"🌐 Localizing CV to {'Hebrew' if lang == 'Hebrew' else 'Arabic'}... (~15 seconds)"):
                    localized = call_claude_localize_cv(english_cv, lang)
                st.session_state[cache_key] = localized
                final_cv = localized
            else:
                final_cv = cached
        else:
            final_cv = english_cv

        st.session_state.final_cv_data = final_cv

        # ── Step 10: Improvement Log ──
        improvement_log = st.session_state.improvement_log
        if improvement_log:
            with st.expander(f"📋 What changed in your CV ({len(improvement_log)} improvements)", expanded=False):
                type_colors = {
                    "Added":       ("#dcfce7", "#166534", "✅"),
                    "Improved":    ("#dbeafe", "#1e40af", "✏️"),
                    "Reorganized": ("#fef9c3", "#713f12", "🔀"),
                }
                for item in improvement_log:
                    itype   = item.get("type", "Improved")
                    section = item.get("section", "")
                    desc    = item.get("description", "")
                    bg, fg, icon = type_colors.get(itype, ("#f1f5f9", "#334155", "•"))
                    st.markdown(
                        f'<div style="background:{bg};border-radius:8px;padding:8px 12px;margin-bottom:6px;font-size:0.88rem;color:{fg};">'
                        f'<strong>{icon} {itype}</strong> — <em>{section}</em>: {desc}</div>',
                        unsafe_allow_html=True
                    )
        elif st.session_state.cv_adjusted:
            # Show summary comparison if no structured log
            original_summary = base_data.get("summary", "")
            adjusted_summary = english_cv.get("summary", "")
            if original_summary and adjusted_summary and original_summary.strip() != adjusted_summary.strip():
                with st.expander("👀 Summary — before vs. after"):
                    c1, c2 = st.columns(2)
                    with c1:
                        st.markdown("**Original:**")
                        st.markdown(f'<div style="background:#fff1f2;padding:12px;border-radius:8px;font-size:0.88rem;color:#334155;">{original_summary}</div>', unsafe_allow_html=True)
                    with c2:
                        st.markdown("**Rewritten:**")
                        st.markdown(f'<div style="background:#f0fdf4;padding:12px;border-radius:8px;font-size:0.88rem;color:#166534;">{adjusted_summary}</div>', unsafe_allow_html=True)

        # ── Regenerate button ──
        col_regen, _ = st.columns([1,4])
        with col_regen:
            if st.button("🔄 Re-generate CV"):
                st.session_state.cv_adjusted = False
                st.session_state.adjusted_cv_data = {}
                st.session_state.final_cv_data = {}
                st.session_state.improvement_log = []
                for k in ["translated_Hebrew", "translated_Arabic"]:
                    if k in st.session_state:
                        del st.session_state[k]
                st.rerun()

        st.markdown("---")

        # ── CV Preview ──
        st.markdown('<div style="background:#f1f5f9;padding:30px 15px;border-radius:12px;">', unsafe_allow_html=True)
        st.markdown(render_cv_html(final_cv, lang), unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)

        st.markdown('<div style="margin-top:1rem;padding:1rem;background:#f0fdf4;border-radius:10px;border:1px solid #bbf7d0;">', unsafe_allow_html=True)
        st.markdown("✅ **Your tailored CV is ready!** Download it in the next tab.")
        st.markdown('</div>', unsafe_allow_html=True)

        nav_buttons(4, can_proceed=True, proceed_label="Next: Download →")


# ============================================================
# TAB 5 — EXPORT
# ============================================================

elif current_tab == 5:
    st.markdown('<div class="rp-page-badge">STEP 6 OF 6</div>', unsafe_allow_html=True)
    st.markdown('<div class="rp-page-title">Download</div>', unsafe_allow_html=True)
    st.markdown('<div class="rp-page-sub">Your tailored resume is ready to send.</div>', unsafe_allow_html=True)

    cv_f = st.session_state.final_cv_data
    has_content = bool(cv_f and any(v.strip() for v in cv_f.values()))

    if not has_content:
        st.warning("Please complete the Preview step first.")
        nav_buttons(5, can_proceed=False)
    else:
        active_lang = st.session_state.active_lang

        # Score badge
        score = st.session_state.analysis_results.get("score", 0)
        if score:
            score_class = "alert-green" if score >= 70 else "alert-orange" if score >= 40 else "alert-red"
            st.markdown(f'<div class="alert-box {score_class}">Match Score: {score}% for {st.session_state.job_role}</div>', unsafe_allow_html=True)

        # File name
        pd_raw = cv_f.get("personal_details","Candidate")
        name_part, _ = format_contact_info(pd_raw)
        first_name = name_part.split()[0] if name_part else "Candidate"
        job_slug = re.sub(r'[^a-zA-Z0-9]', '_', st.session_state.job_role) if st.session_state.job_role else "CV"
        f_name = f"{first_name}_{job_slug}_CV"

        st.markdown("<hr style='border:none;border-top:1px solid #E5E7EB;margin:20px 0;'>", unsafe_allow_html=True)
        c1, c2 = st.columns(2)

        with c1:
            st.markdown('<div class="rp-dl-fmt">DOCX</div><div class="rp-dl-desc">Recommended — send directly to employers</div>', unsafe_allow_html=True)
            docx_buf = generate_docx(cv_f, active_lang)
            st.download_button(
                "📁 Download DOCX",
                docx_buf,
                file_name=f"{f_name}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                type="primary"
            )

        with c2:
            st.markdown('<div class="rp-dl-fmt">TXT</div><div class="rp-dl-desc">For online forms and portals</div>', unsafe_allow_html=True)
            h_map = HEADINGS_MAP[active_lang]
            section_order = [
                ("personal_details", h_map[0]),
                ("summary",          h_map[1]),
                ("education",        h_map[2]),
                ("experience",       h_map[3]),
                ("projects",         h_map[4]),
                ("courses_training", h_map[5]),
                ("volunteering",     h_map[6]),
                ("languages",        h_map[7]),
                ("skills",           h_map[8]),
            ]
            txt_out = f"{st.session_state.job_role} — CV\n{'='*50}\n\n"
            for key, title in section_order:
                v = cv_f.get(key,"").strip()
                if v:
                    if key == "skills":
                        v = re.sub(r'(Technical Skills|Business & Professional Skills|Professional Skills)\s*[,\n]?', '', v, flags=re.IGNORECASE)
                    v = re.sub(r'\s*•\s*', '\n  • ', v)
                    txt_out += f"{title.upper()}\n{'-'*len(title)}\n{v.strip()}\n\n"
            st.download_button(
                "📄 Download TXT",
                txt_out,
                file_name=f"{f_name}.txt",
                mime="text/plain",
                use_container_width=True,
                type="secondary"
            )

        st.markdown("""
        <div class="rp-checklist" style="margin-top:20px;">
            <div class="rp-checklist-title">Before you send</div>
            <div class="rp-check-item">Review every line — make sure everything is accurate.</div>
            <div class="rp-check-item">Never include anything you cannot back up in an interview.</div>
            <div class="rp-check-item">Customize the summary for each job you apply to.</div>
            <div class="rp-check-item">Good luck. You've got this. 🍀</div>
        </div>
        """, unsafe_allow_html=True)
        nav_buttons(5, can_proceed=False)
