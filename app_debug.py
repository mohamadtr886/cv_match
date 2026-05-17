import streamlit as st
import pandas as pd
from io import BytesIO
import docx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
import PyPDF2
import re
from collections import Counter

# --- PAGE CONFIGURATION ---
st.set_page_config(
    page_title="CV Match AI",
    page_icon="💼",
    layout="wide",
    initial_sidebar_state="expanded",
)

# --- INITIALIZE SESSION STATE ---
if 'active_tab' not in st.session_state:
    st.session_state.active_tab = "The Job"
if 'job_role' not in st.session_state:
    st.session_state.job_role = ""
if 'cv_full_text' not in st.session_state:
    st.session_state.cv_full_text = ""
if 'job_desc' not in st.session_state:
    st.session_state.job_desc = ""
if 'input_method' not in st.session_state:
    st.session_state.input_method = "Upload File"
if 'manual_cv_data' not in st.session_state:
    st.session_state.manual_cv_data = {}
if 'follow_up_answers' not in st.session_state:
    st.session_state.follow_up_answers = {}
if 'analysis_results' not in st.session_state:
    st.session_state.analysis_results = {"matches": [], "missing": [], "score": 0}
if 'final_cv_data' not in st.session_state:
    st.session_state.final_cv_data = {}
if 'active_lang' not in st.session_state:
    st.session_state.active_lang = "English"

# --- CUSTOM CSS ---
st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Plus+Jakarta+Sans:wght@400;500;600;700;800&display=swap');
    @import url('https://fonts.googleapis.com/css2?family=EB+Garamond:wght@400;700&family=Inter:wght@400;600;700&display=swap');
    
    html, body, [class*="css"] {
        font-family: 'Plus Jakarta Sans', 'Inter', sans-serif;
    }

    :root {
        --brand-primary: #2563eb;
        --brand-dark: #1e293b;
        --brand-slate: #64748b;
        --bg-main: #f8fafc;
        --white: #ffffff;
        --divider: #e2e8f0;
    }

    .main { background-color: var(--bg-main); }

    .main-header {
        color: var(--brand-dark);
        font-weight: 800;
        font-size: 2.5rem;
        margin-bottom: 0.1rem;
        letter-spacing: -0.04em;
    }
    
    .sub-header {
        color: var(--brand-slate);
        font-size: 1.1rem;
        margin-bottom: 2rem;
        font-weight: 500;
    }

    /* Executive Resume Preview */
    .cv-preview-container {
        background-color: #f1f5f9;
        padding: 40px 20px;
        display: flex;
        justify-content: center;
    }

    .cv-page {
        background-color: white;
        width: 100%;
        max-width: 800px;
        padding: 50px 65px;
        box-shadow: 0 25px 50px -12px rgba(0, 0, 0, 0.15);
        color: #111827;
        line-height: 1.5;
        font-family: 'Inter', sans-serif;
    }

    .cv-header-block {
        text-align: center;
        margin-bottom: 25px;
    }

    .cv-name {
        font-size: 2.25rem;
        font-weight: 800;
        color: #0f172a;
        margin-bottom: 10px;
        text-transform: uppercase;
        letter-spacing: 1px;
    }

    .cv-contact {
        font-size: 0.9rem;
        color: #475569;
        font-weight: 500;
    }

    .cv-header-divider {
        border: none;
        height: 2px;
        background-color: #1e293b;
        margin: 15px 0 25px 0;
    }

    .cv-section-title {
        font-size: 1rem;
        font-weight: 700;
        color: #1e293b;
        text-transform: uppercase;
        letter-spacing: 1.5px;
        margin-top: 25px;
        margin-bottom: 5px;
        border-bottom: 1px solid #cbd5e1;
        padding-bottom: 3px;
    }

    .cv-entry {
        margin-bottom: 15px;
    }

    .cv-entry-header {
        display: flex;
        justify-content: space-between;
        align-items: baseline;
        margin-bottom: 2px;
    }

    .cv-entry-title {
        font-weight: 700;
        font-size: 1.05rem;
        color: #0f172a;
    }

    .cv-entry-date {
        font-weight: 600;
        font-size: 0.85rem;
        color: #64748b;
    }

    .cv-entry-sub {
        font-weight: 600;
        font-size: 0.9rem;
        color: #475569;
        margin-bottom: 6px;
    }

    .cv-bullet-list {
        margin: 0;
        padding-left: 18px;
    }

    .cv-bullet {
        font-size: 0.92rem;
        margin-bottom: 4px;
        color: #334155;
    }

    .cv-skills-grid {
        display: grid;
        grid-template-columns: 1fr 1fr;
        gap: 15px;
        margin-top: 10px;
    }

    .cv-skill-category {
        font-size: 0.9rem;
    }

    .cv-skill-label {
        font-weight: 700;
        color: #1e293b;
        margin-right: 5px;
    }

    .rtl { direction: rtl; text-align: right; }
    .rtl .cv-header-block { text-align: center; }
    .rtl .cv-entry-header { flex-direction: row-reverse; }
    .rtl .cv-bullet-list { padding-right: 18px; padding-left: 0; }

    /* SaaS Navigation Tabs */
    .stTabs [data-baseweb="tab-list"] {
        gap: 8px;
        background-color: transparent;
        padding: 0;
    }

    .stTabs [data-baseweb="tab"] {
        height: 44px;
        white-space: pre;
        background-color: var(--white);
        border: 1px solid var(--slate-200);
        border-radius: 10px;
        padding: 0 20px;
        color: var(--slate-600);
        font-weight: 600;
        transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
    }

    .stTabs [data-baseweb="tab"]:hover {
        border-color: var(--brand-primary);
        color: var(--brand-primary);
    }

    .stTabs [aria-selected="true"] {
        background-color: var(--brand-primary) !important;
        color: var(--white) !important;
        border-color: var(--brand-primary) !important;
        box-shadow: 0 4px 12px rgba(37, 99, 235, 0.2);
    }

    .content-card {
        background-color: var(--white);
        padding: 2.5rem;
        border-radius: 20px;
        border: 1px solid rgba(226, 232, 240, 0.8);
        box-shadow: 0 10px 15px -3px rgba(0, 0, 0, 0.04);
        margin-bottom: 2rem;
    }

    .section-label {
        font-weight: 700;
        color: var(--slate-900);
        margin-bottom: 0.5rem;
        display: block;
        font-size: 1.2rem;
    }

    .keyword-tag {
        display: inline-block;
        background-color: #eff6ff;
        color: #1e40af;
        padding: 4px 12px;
        border-radius: 9999px;
        margin: 4px;
        font-size: 0.85rem;
        font-weight: 600;
        border: 1px solid #dbeafe;
    }

    .stButton > button {
        border-radius: 12px;
        font-weight: 700;
        padding: 0.75rem 2rem;
        background-color: var(--white);
        border: 1px solid var(--slate-200);
        transition: all 0.2s ease;
    }

    .stButton > button:hover {
        border-color: var(--brand-primary);
        color: var(--brand-primary);
        transform: translateY(-2px);
        box-shadow: 0 4px 6px -1px rgba(0,0,0,0.05);
    }
    </style>
    """, unsafe_allow_html=True)

# --- CONSTANTS & DICTIONARIES ---

STOPWORDS = {"a", "an", "the", "and", "or", "but", "if", "then", "else", "when", "at", "from", "by", "for", "with", "about", "against", "between", "into", "through", "during", "before", "after", "above", "below", "to", "from", "up", "down", "in", "out", "on", "off", "over", "under", "again", "further", "then", "once", "here", "there", "all", "any", "both", "each", "few", "more", "most", "other", "some", "such", "no", "nor", "not", "only", "own", "same", "so", "than", "too", "very", "s", "t", "can", "will", "just", "don", "should", "now", "i", "me", "my", "myself", "we", "our", "ours", "ourselves", "you", "your", "yours", "yourself", "yourselves", "he", "him", "his", "himself", "she", "her", "hers", "herself", "it", "its", "itself", "they", "them", "their", "theirs", "themselves", "what", "which", "who", "whom", "this", "that", "these", "those", "am", "is", "are", "was", "were", "be", "been", "being", "have", "has", "had", "having", "do", "does", "did", "doing"}

GENERIC_TERMS = {"ability", "advanced", "bachelor", "common", "core", "critical", "degree", "entry", "familiarity", "field", "fields", "foundation", "job", "knowledge", "large", "like", "strong", "excellent", "good", "responsible", "requirement", "requirements", "responsibilities", "candidate", "company", "team", "work", "working", "looking", "opportunity", "motivated", "passion", "years", "experience", "role", "position", "plus", "must", "needed", "required", "preferred"}

SKILL_KEYWORDS = {
    "python", "sql", "java", "javascript", "react", "node", "html", "css", "git", "github", 
    "r", "matlab", "tableau", "powerbi", "power bi", "excel", "bigquery", "pandas", "numpy", 
    "scikit-learn", "tensorflow", "pytorch", "aws", "azure", "docker", "kubernetes", "unix", 
    "linux", "bash", "spark", "hadoop", "c++", "c#", "php", "ruby", "swift", "kotlin", "typescript"
}

SEMANTIC_MAP = {
    "sql": ["postgresql", "mysql", "databases", "querying", "data manipulation", "relational databases"],
    "python": ["programming", "scripts", "pandas", "numpy", "automation"],
    "react": ["frontend", "javascript", "ui development", "components", "web apps"],
    "data visualization": ["dashboards", "tableau", "powerbi", "plotting", "charts", "power bi"],
    "machine learning": ["ai", "predictive modeling", "scikit-learn", "data science"],
    "version control": ["git", "github", "gitlab", "bitbucket", "svn"],
    "agile": ["scrum", "sprints", "jira", "project management", "kanban"]
}

ACTION_VERBS = {
    "assisted": "Engineered",
    "helped": "Optimized",
    "did": "Developed",
    "worked on": "Spearheaded",
    "responsible for": "Directed",
    "learned": "Mastered",
    "managed": "Architected",
    "wrote": "Formulated"
}

HEADINGS_MAP = {
    "English": ["Personal Details", "Professional Summary", "Education", "Professional Experience", "Projects", "Courses & Training", "Volunteering / Community", "Languages", "Skills", "Additional Information"],
    "Hebrew": ["פרטים אישיים", "תמצית מקצועית", "השכלה", "ניסיון תעסוקתי", "פרויקטים", "קורסים והכשרות", "התנדבות / קהילה", "שפות", "כישורים", "מידע נוסף"],
    "Arabic": ["البيانات الشخصية", "الملخص المهني", "التعليم", "الخبرة العملية", "المشاريع", "الدورات والتدريب", "التطوع / المجتمع", "اللغات", "المهارات", "معلومات إضافية"]
}

CV_TITLES = {"English": "Curriculum Vitae", "Hebrew": "קורות חיים", "Arabic": "السيرة الذاتية"}

SECTION_REGEX = {
    "personal_details": r"(personal details|contact|contact info|פרטים אישיים|البيانات الشخصية)",
    "summary": r"(summary|profile|about me|professional summary|תמצית|פרופיל|תמצית מקצועית|الملخص|ملخص مهني)",
    "education": r"(education|academic|degrees|השכלה|לימודים|التعليم|المؤهلات العلمية)",
    "experience": r"(experience|work|employment|history|professional experience|ניסיון|ניסיון תעסוקתי|الخبرة|الخبرة العملية)",
    "projects": r"(projects|key projects|academic projects|פרויקטים|פרוייקטים|المشاريع)",
    "courses_training": r"(courses|training|certification|certifications|קורסים|הכשרות|הסמכות|الدورات|التدريب)",
    "volunteering": r"(volunteering|community|extracurricular|volunteer|התנדבות|פעילות קהילתית|التطوع)",
    "languages": r"(languages|שפות|اللغات)",
    "skills": r"(skills|competencies|technologies|technical skills|כישורים|מיומנויות|יכולות|المهارات)"
}

SAMPLE_CV_DATA = {
    "personal_details": "Jane Doe | jane.doe@university.edu | +1 555 0123 | LinkedIn: janedoe-profile",
    "summary": "Third-year Computer Science student with strong foundations in software engineering and data analysis. Dedicated to building efficient, scalable solutions and eager to contribute to innovative projects.",
    "education": "B.Sc. in Computer Science, State University (Expected May 2025)\nRelevant Coursework: Data Structures, Algorithms, Database Systems, Web Development",
    "experience": "Summer Research Intern, AI Lab (2023)\n- Assisted in developing data preprocessing scripts in Python.\n- Conducted performance benchmarking for machine learning models.\n- Documented experimental results for team review.",
    "projects": "OpenSource Contributor - Improved UI components for a React-based dashboard project.\nWeather Insight App - Built a responsive weather tracking application using OpenWeather API.",
    "courses_training": "AWS Cloud Practitioner Essentials, Advanced SQL Certification (Coursera)",
    "volunteering": "Peer Tutor - Provided technical assistance to first-year CS students in Java and C++.",
    "languages": "English (Native), French (Conversational)",
    "skills": "Python, Java, SQL, React, Git, Docker, Problem Solving, Technical Writing"
}

# --- CORE FUNCTIONS ---

def detect_language(text):
    if re.search(r'[\u0590-\u05FF]', text): return "Hebrew"
    if re.search(r'[\u0600-\u06FF]', text): return "Arabic"
    return "English"

def clean_text_simple(text):
    text = text.lower()
    text = re.sub(r'[^a-z0-9\u0590-\u05FF\u0600-\u06FF\s]', ' ', text)
    return text

def extract_keywords(text):
    cleaned = clean_text_simple(text)
    words = cleaned.split()
    return [w for w in words if w not in STOPWORDS and w not in GENERIC_TERMS and len(w) > 2]

def extract_text_from_pdf(file):
    try:
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    except: return ""

def extract_text_from_docx(file):
    try:
        doc = docx.Document(file)
        return "\n".join([para.text for para in doc.paragraphs])
    except: return ""

def classify_sections_refined(text):
    sections = {"personal_details": "", "summary": "", "education": "", "experience": "", "projects": "", "courses_training": "", "volunteering": "", "languages": "", "skills": "", "additional_information": ""}
    current_section = "personal_details"
    lines = text.split('\n')
    for line in lines:
        clean_line = line.strip()
        if not clean_line: continue
        low_line = clean_line.lower()
        found_header = False
        if len(clean_line) < 35:
            for key, pattern in SECTION_REGEX.items():
                if re.search(pattern, low_line):
                    current_section = key
                    found_header = True
                    break
        if not found_header:
            if current_section in ["experience", "education", "personal_details"] and len(clean_line) < 60:
                line_words = set(re.findall(r'\w+', low_line))
                if line_words.intersection(SKILL_KEYWORDS) and len(line_words) < 8:
                    sections["skills"] += clean_line + ", "
                    continue
            sections[current_section] += clean_line + "\n"
    for k in sections:
        sections[k] = re.sub(r'\n{2,}', '\n', sections[k].strip())
        if k == "skills":
            sections[k] = sections[k].strip(", ")
    return sections

def format_contact_info(text):
    lines = text.split('\n')
    if not lines: return "", ""
    name = lines[0].strip()
    rest = " | ".join([l.strip() for l in lines[1:] if l.strip()])
    return name, rest

def semantic_match(cv_keywords, job_keywords):
    matches, missing = [], []
    cv_set, job_set = set(cv_keywords), set(job_keywords)
    for jk in job_set:
        found = False
        if jk in cv_set:
            matches.append(jk)
            found = True
        else:
            for primary, synonyms in SEMANTIC_MAP.items():
                if jk == primary or jk in synonyms:
                    if any(s in cv_set for s in synonyms) or primary in cv_set:
                        matches.append(jk)
                        found = True
                        break
        if not found: missing.append(jk)
    return matches, missing

def get_ai_suggestion(role, field):
    suggestions = {
        "summary": f"Ambitious {role} student focused on leveraging technical skills to build innovative solutions.",
        "experience": f"Collaborated with cross-functional teams to deliver {role}-related projects.",
        "projects": f"Developed a scalable application that demonstrates core {role} principles."
    }
    return suggestions.get(field, "Start typing your experience here...")

# --- CORE FUNCTIONS (UPGRADED INTEGRATION & RENDERING) ---

def group_skills_segregated(skills_text):
    if not skills_text: return {"PROFESSIONAL": [], "TECHNICAL": []}
    skills = [s.strip() for s in re.split(r'[,|\n]', skills_text) if s.strip()]
    tech_keywords = set(list(SKILL_KEYWORDS) + ["sql", "python", "aws", "docker", "git", "react", "tableau", "powerbi", "excel", "data analysis"])
    technical, professional = [], []
    for s in skills:
        sl = s.lower()
        if any(kw in sl for kw in tech_keywords) or len(s) < 15: technical.append(s)
        else: professional.append(s)
    if not professional and technical:
        professional = [s for s in technical if len(s) > 10][:3]
        technical = [s for s in technical if s not in professional]
    return {"PROFESSIONAL": list(set(professional)), "TECHNICAL": list(set(technical))}

def integrate_refinement_smart(cv_data, answers, job_role, missing_keywords):
    upgraded_cv = cv_data.copy()
    def upgrade_text(text):
        if not text: return ""
        for weak, strong in ACTION_VERBS.items():
            text = re.sub(rf'\b{weak}\b', strong, text, flags=re.IGNORECASE)
        return text
    for sec in ["summary", "experience", "projects", "education"]: upgraded_cv[sec] = upgrade_text(upgraded_cv.get(sec, ""))
    answers_text = " ".join(answers.values()).lower()
    current_skills = upgraded_cv.get("skills", "").lower()
    new_skills = [mk for mk in missing_keywords if mk.lower() in answers_text and mk.lower() not in current_skills]
    if new_skills: upgraded_cv["skills"] = upgraded_cv.get("skills", "") + ", " + ", ".join(new_skills)
    achievement = ""
    for ans in answers.values():
        if re.search(r'\d+%|\d+x|improved|optimized|spearheaded', ans.lower()):
            achievement = ans.strip()
            break
    if achievement and upgraded_cv.get("experience"):
        exp_lines = upgraded_cv["experience"].split('\n')
        for i, line in enumerate(exp_lines):
            if line.strip().startswith(('-', '•', '*')):
                exp_lines[i] = f"* {achievement}"
                break
        upgraded_cv["experience"] = "\n".join(exp_lines)
    if job_role and upgraded_cv.get("summary"):
        if job_role.lower() not in upgraded_cv["summary"].lower():
            upgraded_cv["summary"] = f"Goal-oriented professional with a focus on **{job_role}**. {upgraded_cv['summary']}"
    return upgraded_cv

def render_executive_cv_html(cv_data, lang="English"):
    is_rtl = lang in ["Hebrew", "Arabic"]
    l_class, dir_attr = ("rtl" if is_rtl else "ltr"), ('dir="rtl"' if is_rtl else 'dir="ltr"')
    h_map = HEADINGS_MAP[lang]
    name_contact = cv_data.get("personal_details", "CANDIDATE NAME | Contact Info")
    name, contact = format_contact_info(name_contact)
    if not name: name = "CANDIDATE NAME"
    html = f'<div class="cv-page {l_class}" {dir_attr}><div class="cv-header-block"><div class="cv-name">{name.upper()}</div><div class="cv-contact">{contact}</div><hr class="cv-header-divider"></div>'
    if cv_data.get("summary"): html += f'<div class="cv-section-title">{h_map[1].upper()}</div><div class="cv-entry" style="font-size:0.95rem; margin-top:8px;">{cv_data["summary"]}</div>'
    if cv_data.get("experience"):
        html += f'<div class="cv-section-title">{h_map[3].upper()}</div>'
        for entry in cv_data["experience"].split('\n\n'):
            lines = [l.strip() for l in entry.split('\n') if l.strip()]
            if not lines: continue
            html += f'<div class="cv-entry"><div class="cv-entry-header"><span class="cv-entry-title">{lines[0]}</span></div><div class="cv-entry-sub">{lines[1] if len(lines) > 1 else ""}</div><ul class="cv-bullet-list">'
            for b in lines[2:]:
                style = 'style="list-style-type: circle; margin-left: 20px; font-size: 0.85rem; color: #475569;"' if b.startswith(('  ', '\t', '-')) else ''
                html += f'<li class="cv-bullet" {style}>{b.strip("- •*")}</li>'
            html += "</ul></div>"
    if cv_data.get("projects"):
        html += f'<div class="cv-section-title">{h_map[4].upper()}</div>'
        for proj in cv_data["projects"].split('\n\n'):
            lines = [l.strip() for l in proj.split('\n') if l.strip()]
            if not lines: continue
            html += f'<div class="cv-entry"><div class="cv-entry-title">{lines[0]}</div><ul class="cv-bullet-list">'
            for b in lines[1:]: html += f'<li class="cv-bullet">{b.strip("- •*")}</li>'
            html += "</ul></div>"
    if cv_data.get("education"):
        html += f'<div class="cv-section-title">{h_map[2].upper()}</div>'
        for edu in cv_data["education"].split('\n\n'):
            lines = [l.strip() for l in edu.split('\n') if l.strip()]
            if not lines: continue
            html += f'<div class="cv-entry"><div class="cv-entry-header"><span class="cv-entry-title">{lines[0]}</span></div><div class="cv-entry-sub">{lines[1] if len(lines) > 1 else ""}</div></div>'
    skills_data = group_skills_segregated(cv_data.get("skills", ""))
    if skills_data["PROFESSIONAL"]: html += f'<div class="cv-section-title">PROFESSIONAL SKILLS</div><div class="cv-entry" style="font-size:0.92rem; margin-top:8px;">{", ".join(skills_data["PROFESSIONAL"])}</div>'
    if skills_data["TECHNICAL"]: html += f'<div class="cv-section-title">TECHNICAL SKILLS</div><div class="cv-entry" style="font-size:0.92rem; margin-top:8px;">{", ".join(skills_data["TECHNICAL"])}</div>'
    html += "</div>"
    return html

def generate_docx_executive(cv_data, output_lang):
    doc = Document()
    for section in doc.sections: section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = docx.shared.Inches(1)
    is_rtl = output_lang in ["Hebrew", "Arabic"]
    h_map = HEADINGS_MAP[output_lang]
    style = doc.styles['Normal']
    style.font.name, style.font.size = ('Arial' if is_rtl else 'Calibri'), Pt(11)
    name, contact = format_contact_info(cv_data.get("personal_details", ""))
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_name = p_name.add_run(name.upper())
    run_name.bold, run_name.font.size = True, Pt(22)
    p_contact = doc.add_paragraph(contact)
    p_contact.alignment, p_contact.paragraph_format.space_after = WD_ALIGN_PARAGRAPH.CENTER, Pt(6)
    doc.add_paragraph("_" * 80).alignment = WD_ALIGN_PARAGRAPH.CENTER
    order = [("summary", 1), ("experience", 3), ("projects", 4), ("education", 2)]
    for key, h_idx in order:
        content = cv_data.get(key, "").strip()
        if content:
            h = doc.add_paragraph()
            if is_rtl: h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            h_run = h.add_run(h_map[h_idx].upper())
            h_run.bold, h_run.font.size = True, Pt(12)
            h.paragraph_format.space_before = Pt(12)
            if key in ["experience", "projects"]:
                for entry in content.split('\n\n'):
                    lines = [l.strip() for l in entry.split('\n') if l.strip()]
                    for i, line in enumerate(lines):
                        p = doc.add_paragraph()
                        if is_rtl: p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        if i == 0: p.add_run(line).bold = True
                        elif i == 1 and key == "experience": p.add_run(line).italic = True
                        else:
                            p_bullet = doc.add_paragraph(style='List Bullet')
                            if is_rtl: p_bullet.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            p_bullet.add_run(line.strip("- •*"))
            else:
                p = doc.add_paragraph(content)
                if is_rtl: p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    skills_data = group_skills_segregated(cv_data.get("skills", ""))
    if skills_data["PROFESSIONAL"]:
        h = doc.add_paragraph()
        if is_rtl: h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        h.add_run("PROFESSIONAL SKILLS").bold = True
        doc.add_paragraph(", ".join(skills_data["PROFESSIONAL"]))
    if skills_data["TECHNICAL"]:
        h = doc.add_paragraph()
        if is_rtl: h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        h.add_run("TECHNICAL SKILLS").bold = True
        doc.add_paragraph(", ".join(skills_data["TECHNICAL"]))
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- APP LAYOUT (TABS) ---
st.markdown('<div class="main-header">CV Match AI</div>', unsafe_allow_html=True)
st.markdown('<div class="sub-header">Premium Career Strategy & Tailoring Suite</div>', unsafe_allow_html=True)
tab_titles = ["The Job", "CV Content", "Analysis", "Refinement", "Preview", "Export"]
tabs = st.tabs(tab_titles)

with tabs[0]:
    st.markdown('<div class="content-card"><span class="section-label">Target Position & Requirements</span>', unsafe_allow_html=True)
    st.write("Define your target. Paste the description or insert a link to start the semantic analysis.")
    st.session_state.job_role = st.text_input("What is your target job title?", value=st.session_state.job_role, placeholder="e.g., Software Engineering Intern")
    input_mode = st.radio("How would you like to provide the job details?", ["Paste Job Text", "Insert Job Link"], horizontal=True)
    if input_mode == "Insert Job Link":
        job_url = st.text_input("Job Posting URL:", placeholder="https://linkedin.com/jobs/...")
        if job_url:
            st.info("🔗 Link detected. (Simulated extraction in progress...)")
            if st.button("Extract from URL"): st.session_state.job_desc = f"Simulated extraction for: {st.session_state.job_role}. Requirements: Python, SQL, Git, Teamwork, Data Analysis."
    else: st.session_state.job_desc = st.text_area("Job Description:", value=st.session_state.job_desc, height=250, placeholder="Paste requirements...")
    if st.button("✨ Load Demo Job Description"):
        st.session_state.job_role, st.session_state.job_desc = "Junior Data Analyst", "Requirements: Proficient in SQL and Python. Experience with Tableau or PowerBI. Knowledge of Git."
        st.rerun()
    st.markdown('</div>', unsafe_allow_html=True)

with tabs[1]:
    st.markdown('<div class="content-card"><span class="section-label">Your Professional Profile</span>', unsafe_allow_html=True)
    st.session_state.input_method = st.radio("Choose input method:", ["Upload File", "Paste Text", "Manual Entry"], horizontal=True)
    if st.session_state.input_method == "Upload File":
        uploaded_file = st.file_uploader("Upload CV", type=["pdf", "docx", "txt"])
        if uploaded_file:
            if uploaded_file.type == "application/pdf": st.session_state.cv_full_text = extract_text_from_pdf(uploaded_file)
            elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document": st.session_state.cv_full_text = extract_text_from_docx(uploaded_file)
            else: st.session_state.cv_full_text = str(uploaded_file.read(), "utf-8")
            st.success("File processed.")
    elif st.session_state.input_method == "Paste Text":
        st.session_state.cv_full_text = st.text_area("Paste CV Text:", value=st.session_state.cv_full_text, height=300)
        if st.button("✨ Auto-Structure Pasted Text"):
            st.session_state.manual_cv_data = classify_sections_refined(st.session_state.cv_full_text)
            st.success("Text structured successfully.")
    elif st.session_state.input_method == "Manual Entry":
        m_fields = {"summary": ("Your Hook / Career Goal", "Hardworking CS student..."), "experience": ("Work or Volunteer Experience", "Internships..."), "projects": ("Academic or Side Projects", "University projects..."), "education": ("Education", "University..."), "skills": ("Skills & Technologies", "Python..."), "personal_details": ("Full Name & Contact Info", "Name | Email...")}
        c1, c2 = st.columns(2)
        with c1:
            for key in ["personal_details", "summary", "education"]:
                st.write(f"**{m_fields[key][0]}**")
                if st.button(f"✨ Suggest {m_fields[key][0]}", key=f"ai_{key}"): st.session_state.manual_cv_data[key] = get_ai_suggestion(st.session_state.job_role, key)
                st.session_state.manual_cv_data[key] = st.text_area(m_fields[key][0], value=st.session_state.manual_cv_data.get(key, ""), placeholder=m_fields[key][1], key=f"input_{key}", label_visibility="collapsed", height=100)
        with c2:
            for key in ["experience", "projects", "skills"]:
                st.write(f"**{m_fields[key][0]}**")
                if st.button(f"✨ Suggest {m_fields[key][0]}", key=f"ai_{key}"): st.session_state.manual_cv_data[key] = get_ai_suggestion(st.session_state.job_role, key)
                st.session_state.manual_cv_data[key] = st.text_area(m_fields[key][0], value=st.session_state.manual_cv_data.get(key, ""), placeholder=m_fields[key][1], key=f"input_{key}", label_visibility="collapsed", height=100)
    st.markdown('</div>', unsafe_allow_html=True)

with tabs[2]:
    cv_p = st.session_state.cv_full_text or any(st.session_state.manual_cv_data.values())
    if not st.session_state.job_desc or not cv_p: st.warning("Please provide Job Description and CV content.")
    else:
        st.markdown('<div class="content-card"><span class="section-label">Semantic Match Analysis</span>', unsafe_allow_html=True)
        cv_text = st.session_state.cv_full_text if st.session_state.cv_full_text else "\n".join(st.session_state.manual_cv_data.values())
        matches, missing = semantic_match(extract_keywords(cv_text), extract_keywords(st.session_state.job_desc))
        score = int((len(matches) / max(len(extract_keywords(st.session_state.job_desc)), 1)) * 100)
        col1, col2 = st.columns([1, 2])
        with col1:
            st.write("### Match Score")
            st.markdown(f"<h1 style='color:var(--brand-primary); font-size:4rem;'>{score}%</h1>", unsafe_allow_html=True)
            st.progress(score / 100)
        with col2:
            st.write("### Strategy")
            if score < 50: st.error("Significant gap detected.")
            elif score < 80: st.warning("Strong foundation.")
            else: st.success("Excellent alignment!")
        st.write("---")
        m_col, g_col = st.columns(2)
        with m_col:
            st.write("**Identified Matches:**")
            st.markdown("".join([f'<span class="keyword-tag">{k}</span>' for k in matches[:15]]), unsafe_allow_html=True)
        with g_col:
            st.write("**Critical Gaps:**")
            for m in missing[:5]: st.markdown(f"<div style='background-color:#fff1f2; border-left:4px solid #e11d48; padding:8px; margin-bottom:8px; border-radius:4px;'><strong>{m.capitalize()}</strong><br><span style='font-size:0.85rem; color:#9f1239;'>Emphasize this skill.</span></div>", unsafe_allow_html=True)
        st.session_state.analysis_results = {"matches": matches, "missing": missing, "score": score}
        st.markdown('</div>', unsafe_allow_html=True)

with tabs[3]:
    st.markdown('<div class="content-card"><span class="section-label">Refinement Questions</span>', unsafe_allow_html=True)
    missing = st.session_state.analysis_results.get("missing", [])
    if not missing: st.write("No major gaps found!")
    else:
        q1 = f"The role requires **{missing[0]}**. Describe a project where you used this tool."
        q2 = f"How have you applied **{missing[1] if len(missing)>1 else 'problem solving'}** to improve an outcome?"
        st.write(f"**Question 1:** {q1}")
        st.session_state.follow_up_answers["q1"] = st.text_area("A1", value=st.session_state.follow_up_answers.get("q1", ""), key="ans1", label_visibility="collapsed")
        st.write(f"**Question 2:** {q2}")
        st.session_state.follow_up_answers["q2"] = st.text_area("A2", value=st.session_state.follow_up_answers.get("q2", ""), key="ans2", label_visibility="collapsed")
    st.markdown('</div>', unsafe_allow_html=True)

with tabs[4]:
    st.markdown('<div class="content-card"><span class="section-label">Executive CV Preview</span>', unsafe_allow_html=True)
    st.session_state.active_lang = st.radio("Choose Document Language:", ["English", "Hebrew", "Arabic"], horizontal=True, key="lang_select")
    base_data = st.session_state.manual_cv_data.copy() if any(st.session_state.manual_cv_data.values()) else classify_sections_refined(st.session_state.cv_full_text)
    final_cv = integrate_refinement_smart(base_data, st.session_state.follow_up_answers, st.session_state.job_role, st.session_state.analysis_results.get("missing", []))
    st.session_state.final_cv_data = final_cv
    st.markdown('<div class="cv-preview-container">', unsafe_allow_html=True)
    st.markdown(render_executive_cv_html(final_cv, st.session_state.active_lang), unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)
    st.success("✨ Ready for download in the next tab.")
    st.markdown('</div>', unsafe_allow_html=True)

with tabs[5]:
    st.markdown('<div class="content-card" style="text-align:center;"><span class="section-label">Final Export</span>', unsafe_allow_html=True)
    if not st.session_state.final_cv_data: st.warning("Please preview your CV first.")
    else:
        cv_f = st.session_state.final_cv_data
        u_name = cv_f["personal_details"].split('|')[0].strip().split()[0] if cv_f.get("personal_details") else "Candidate"
        f_name = f"{u_name}_{st.session_state.job_role.replace(' ', '_')}_CV"
        c1, c2 = st.columns(2)
        with c1:
            txt_out = f"{f_name.upper()}\n{'='*len(f_name)}\n"
            for k, v in cv_f.items():
                if v.strip(): txt_out += f"\n{k.upper()}\n{v}\n"
            st.download_button("📄 Download TXT", txt_out, file_name=f"{f_name}.txt", use_container_width=True)
        with c2:
            st.download_button("📁 Download DOCX", generate_docx_executive(cv_f, st.session_state.active_lang), file_name=f"{f_name}.docx", use_container_width=True)
    st.markdown('</div>', unsafe_allow_html=True)

