import streamlit as st
import pandas as pd
from io import BytesIO
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import PyPDF2
import time
import re

# --- PAGE CONFIGURATION ---
st.set_page_config(
    page_title="Universal CV Match AI",
    page_icon="💼",
    layout="wide",
    initial_sidebar_state="expanded",
)

# --- CONSTANTS & DATA ---
SECTION_HEADINGS = {
    "en": {
        "personal": "Personal Details",
        "summary": "Professional Summary",
        "education": "Education",
        "experience": "Professional Experience",
        "courses": "Courses & Training",
        "volunteering": "Volunteering / Community",
        "languages": "Languages",
        "skills": "Skills"
    },
    "he": {
        "personal": "פרטים אישיים",
        "summary": "תמצית",
        "education": "השכלה",
        "experience": "ניסיון תעסוקתי",
        "courses": "קורסים והכשרות",
        "volunteering": "התנדבות / פעילות קהילתית",
        "languages": "שפות",
        "skills": "כישורים"
    }
}

SECTION_KEYWORDS = {
    "en": {
        "summary": ["summary", "profile", "professional summary", "about", "objective"],
        "education": ["education", "academic background", "studies"],
        "experience": ["experience", "work experience", "professional experience", "employment history", "employment"],
        "projects": ["projects", "academic projects", "key projects"],
        "skills": ["skills", "technical skills", "competencies", "technologies"],
        "languages": ["languages"],
        "courses": ["certifications", "courses", "training", "professional development", "certificates"],
        "volunteering": ["volunteering", "leadership", "community", "social activity"],
        "contact": ["contact", "personal details", "personal information", "profile information"]
    },
    "he": {
        "summary": ["תמצית", "פרופיל אישי", "תקציר", "תמצית מקצועית"],
        "education": ["השכלה", "רקע אקדמי", "לימודים"],
        "experience": ["ניסיון", "ניסיון תעסוקתי", "ניסיון מקצועי", "ניסיון רלוונטי", "תעסוקה"],
        "projects": ["פרויקטים", "פרויקטים אקדמיים"],
        "skills": ["כישורים", "מיומנויות", "כישורים טכניים", "מיומנות"],
        "tools": ["כלים", "טכנולוגיות"],
        "languages": ["שפות"],
        "courses": ["קורסים", "הכשרות", "תעודות"],
        "volunteering": ["התנדבות", "מנהיגות", "פעילות קהילתית"],
        "contact": ["פרטים אישיים", "יצירת קשר", "פרטי קשר"]
    }
}

SKILLS_KEYWORDS = [
    "Python", "SQL", "R", "Java", "JavaScript", "Excel", "Power BI", "Tableau",
    "Machine Learning", "Statistics", "Project Management", "Agile", "Scrum",
    "Communication", "Teamwork", "Analytical", "Problem Solving", "Leadership",
    "Customer Service", "Marketing", "Sales", "React", "Node.js", "AWS", "Git"
]

# --- SESSION STATE ---
if 'step' not in st.session_state: st.session_state.step = 1
if 'cv_raw' not in st.session_state: st.session_state.cv_raw = ""
if 'cv_sections' not in st.session_state: 
    st.session_state.cv_sections = {k: "" for k in SECTION_HEADINGS["en"].keys()}
if 'job_desc' not in st.session_state: st.session_state.job_desc = ""
if 'job_role' not in st.session_state: st.session_state.job_role = ""
if 'lang' not in st.session_state: st.session_state.lang = "en"
if 'direction' not in st.session_state: st.session_state.direction = "ltr"
if 'match_data' not in st.session_state: st.session_state.match_data = {}
if 'follow_up' not in st.session_state: st.session_state.follow_up = {}
if 'extra_info' not in st.session_state: st.session_state.extra_info = ""

# --- CORE FUNCTIONS ---

def extract_text_from_file(file):
    try:
        if file.name.lower().endswith(".pdf"):
            reader = PyPDF2.PdfReader(file)
            return "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
        elif file.name.lower().endswith(".docx"):
            doc = docx.Document(file)
            return "\n".join([p.text for p in doc.paragraphs])
        elif file.name.lower().endswith(".txt"):
            return str(file.read(), "utf-8")
        else:
            st.error(f"Unsupported file format: {file.name}")
            return ""
    except Exception as e:
        st.error(f"Error extracting text from {file.name}: {e}")
        return ""

def clean_cv_text(text):
    if not text: return ""
    text = re.sub(r'\n\s*\n', '\n', text)
    lines = text.split('\n')
    cleaned_lines = []
    for line in lines:
        line = line.strip()
        if len(line.split()) < 3 and not re.search(r'[@\d]', line): continue
        line = " ".join(line.split())
        line = re.sub(r'^[•\-\*\.]\s*', '', line)
        if line: cleaned_lines.append(line)
    return "\n".join(cleaned_lines)

def detect_language_and_direction(text):
    if re.search(r'[\u0590-\u05FF]', text): return "he", "rtl"
    return "en", "ltr"

def parse_cv_sections(text, lang):
    lines = text.split('\n')
    parsed = {k: [] for k in SECTION_HEADINGS[lang].keys()}
    parsed["unclassified"] = []
    current_section = "personal"
    keywords_map = SECTION_KEYWORDS[lang]
    for line in lines:
        clean_line = line.strip().lower().replace(":", "")
        found_header = False
        for sec_name, keywords in keywords_map.items():
            if clean_line in [k.lower() for k in keywords]:
                current_section = "personal" if sec_name == "contact" else sec_name
                found_header = True
                break
        if not found_header:
            target = current_section if current_section in parsed else "unclassified"
            parsed[target].append(line.strip())
    return {k: clean_cv_text("\n".join(v)) for k, v in parsed.items()}

def extract_job_requirements(job_desc):
    job_desc_lower = job_desc.lower()
    return list(set([kw for kw in SKILLS_KEYWORDS if kw.lower() in job_desc_lower]))

def compare_cv_to_job(cv_text, job_reqs):
    cv_text_lower = cv_text.lower()
    matching = [req for req in job_reqs if req.lower() in cv_text_lower]
    missing = [req for req in job_reqs if req.lower() not in cv_text_lower]
    score = (len(matching) / len(job_reqs) * 100) if job_reqs else 0
    return {"score": round(score), "matching": matching, "missing": missing}

def create_docx(cv_data, lang, direction):
    doc = docx.Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial' if direction == "rtl" else 'Calibri'
    font.size = Pt(11)
    personal = cv_data.get("personal", "").split('\n')
    name = personal[0] if personal else "Curriculum Vitae"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(name)
    run.bold, run.font.size = True, Pt(16)
    if len(personal) > 1:
        c = doc.add_paragraph(" | ".join(personal[1:]))
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titles, order = SECTION_HEADINGS[lang], ["summary", "education", "experience", "courses", "volunteering", "languages", "skills"]
    for sec in order:
        content = cv_data.get(sec, "").strip()
        if content:
            h = doc.add_paragraph()
            h.alignment = WD_ALIGN_PARAGRAPH.RIGHT if direction == "rtl" else WD_ALIGN_PARAGRAPH.LEFT
            h_run = h.add_run(titles[sec])
            h_run.bold, h_run.font.size, h_run.underline = True, Pt(12), True
            lines = content.split('\n')
            for line in lines:
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT if direction == "rtl" else WD_ALIGN_PARAGRAPH.LEFT
                if sec in ["experience", "courses", "volunteering", "skills"] and len(lines) > 1:
                    para.style = 'List Bullet'
                    para.add_run(line)
                else: para.add_run(line)
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- UI HELPERS ---

def render_cv_preview(cv_data, lang, direction):
    align = "right" if direction == "rtl" else "left"
    st.markdown(f"""
        <style>
        .cv-page {{ background-color: white; padding: 50px; border: 1px solid #ddd; border-radius: 2px; box-shadow: 0 0 10px rgba(0,0,0,0.1); color: #333; font-family: 'Arial', sans-serif; direction: {direction}; text-align: {align}; max-width: 800px; margin: auto; }}
        .cv-header {{ text-align: center; border-bottom: 2px solid #1e3a8a; margin-bottom: 25px; padding-bottom: 15px; }}
        .cv-name {{ font-size: 24pt; font-weight: bold; color: #1e3a8a; margin: 0; }}
        .cv-contact {{ font-size: 10pt; color: #555; margin-top: 5px; }}
        .section-title {{ font-size: 14pt; font-weight: bold; color: #1e3a8a; border-bottom: 1px solid #ccc; margin-top: 20px; margin-bottom: 10px; text-transform: uppercase; }}
        .section-content {{ font-size: 11pt; line-height: 1.5; margin-bottom: 15px; }}
        .bullet-item {{ display: list-item; margin-left: { "0" if direction == "rtl" else "20px" }; margin-right: { "20px" if direction == "rtl" else "0" }; list-style-type: disc; }}
        </style>
    """, unsafe_allow_html=True)
    with st.container():
        st.markdown('<div class="cv-page">', unsafe_allow_html=True)
        personal_lines = cv_data.get("personal", "").split('\n')
        name = personal_lines[0] if personal_lines else "CV"
        st.markdown(f'<div class="cv-header"><div class="cv-name">{name}</div>', unsafe_allow_html=True)
        if len(personal_lines) > 1: st.markdown(f'<div class="cv-contact">{" | ".join(personal_lines[1:])}</div>', unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)
        titles, order = SECTION_HEADINGS[lang], ["summary", "education", "experience", "courses", "volunteering", "languages", "skills"]
        for sec in order:
            content = cv_data.get(sec, "").strip()
            if content:
                st.markdown(f'<div class="section-title">{titles[sec]}</div>', unsafe_allow_html=True)
                lines = content.split('\n')
                st.markdown('<div class="section-content">', unsafe_allow_html=True)
                for line in lines:
                    if sec in ["experience", "courses", "volunteering", "skills"] and len(lines) > 1: st.markdown(f'<div class="bullet-item">{line}</div>', unsafe_allow_html=True)
                    else: st.markdown(f'<div>{line}</div>', unsafe_allow_html=True)
                st.markdown('</div>', unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)

# --- APP LAYOUT ---

st.markdown("<h1 style='text-align: center; color: #1e3a8a;'>Universal CV Match AI 🚀</h1>", unsafe_allow_html=True)

# SIDEBAR
with st.sidebar:
    st.image("https://cdn-icons-png.flaticon.com/512/3135/3135673.png", width=80)
    st.title("Settings")
    if st.button("Restart Process 🔄"):
        for key in list(st.session_state.keys()): del st.session_state[key]
        st.rerun()

# STEP 1: INPUT
if st.session_state.step == 1:
    col1, col2 = st.columns([1, 1])
    with col1:
        st.markdown("### 1. Your Information")
        mode = st.radio("Choose Input Method:", [
            "Upload current / initial CV with all information needed for tailoring",
            "Paste current / initial CV text",
            "Build CV manually from my information",
            "Use sample CV for testing"
        ], horizontal=False)
        
        if "Upload" in mode:
            st.markdown("##### Upload your current or initial CV")
            st.caption("The file should include your education, experience, projects, courses, skills, languages, and any information you want the app to use. You can add missing details below.")
            f = st.file_uploader("Upload (PDF, DOCX, TXT)", type=["pdf", "docx", "txt"])
            if f:
                st.session_state.cv_raw = extract_text_from_file(f)
                if st.session_state.cv_raw:
                    st.session_state.lang, st.session_state.direction = detect_language_and_direction(st.session_state.cv_raw)
                    st.session_state.cv_sections = parse_cv_sections(st.session_state.cv_raw, st.session_state.lang)
                    st.success("File processed!")
            st.session_state.extra_info = st.text_area("Add anything important that is not clearly written in your CV", help="Optional extra information")
        
        elif "Paste" in mode:
            st.session_state.cv_raw = st.text_area("Paste your CV here:", height=250)
            if st.button("Process Text"):
                st.session_state.lang, st.session_state.direction = detect_language_and_direction(st.session_state.cv_raw)
                st.session_state.cv_sections = parse_cv_sections(st.session_state.cv_raw, st.session_state.lang)
                st.success("Text processed!")

        elif "Build CV manually" in mode:
            with st.expander("Structured CV Sections", expanded=True):
                for key in st.session_state.cv_sections.keys():
                    if key != "unclassified":
                        st.session_state.cv_sections[key] = st.text_area(f"{key.capitalize()}", st.session_state.cv_sections[key], height=100)

        elif "Use sample CV" in mode:
            if st.button("Load Sample Data"):
                st.session_state.cv_raw = "Alex Smith\nalex@example.com | 050-1234567\nEducation: B.Sc. Computer Science, TAU\nExperience: Python Developer Intern at TechCo. Built data pipelines.\nSkills: Python, SQL, Git."
                st.session_state.lang, st.session_state.direction = detect_language_and_direction(st.session_state.cv_raw)
                st.session_state.cv_sections = parse_cv_sections(st.session_state.cv_raw, st.session_state.lang)

    with col2:
        st.markdown("### 2. Target Job")
        st.session_state.job_role = st.text_input("Job Role Name", st.session_state.job_role)
        st.session_state.job_desc = st.text_area("Job Description / Requirements", st.session_state.job_desc, height=300)

    if st.button("Generate Tailored CV ➡️", use_container_width=True):
        if len(st.session_state.job_desc.split()) < 30 and st.session_state.job_desc.lower() != "whatever":
            st.error("Job description must be at least 30 words.")
        elif not any(st.session_state.cv_sections.values()):
            st.error("Please provide your CV info.")
        else:
            if st.session_state.extra_info:
                st.session_state.cv_sections["experience"] += f"\n{st.session_state.extra_info}"
            all_text = "\n".join(st.session_state.cv_sections.values())
            reqs = extract_job_requirements(st.session_state.job_desc)
            st.session_state.match_data = compare_cv_to_job(all_text, reqs)
            st.session_state.step = 2
            st.rerun()

# STEP 2: TAILORING
elif st.session_state.step == 2:
    st.markdown(f"## Tailoring for {st.session_state.job_role}")
    st.metric("Match Score", f"{st.session_state.match_data['score']}%")
    st.markdown("### 💡 Quick Questions to Improve Your Score")
    with st.form("tailor_form"):
        for skill in st.session_state.match_data['missing'][:5]:
            st.session_state.follow_up[skill] = st.text_input(f"Do you have experience with **{skill}**? Describe briefly:")
        if st.form_submit_button("Generate Professional Final CV ✨"):
            for skill, ans in st.session_state.follow_up.items():
                if ans.strip():
                    st.session_state.cv_sections["experience"] += f"\nUsed {skill} to {ans.strip()}"
                    if skill not in st.session_state.cv_sections["skills"]: st.session_state.cv_sections["skills"] += f", {skill}"
            st.session_state.step = 3
            st.rerun()

# STEP 3: FINAL REVIEW & DOWNLOAD
elif st.session_state.step == 3:
    st.markdown("## 🏁 Step 3: Review & Download")
    with st.expander("📝 Review and Edit Final CV Content Before Download", expanded=False):
        for key in st.session_state.cv_sections.keys():
            if key != "unclassified":
                st.session_state.cv_sections[key] = st.text_area(f"Final {key.capitalize()}", st.session_state.cv_sections[key], height=150)
    render_cv_preview(st.session_state.cv_sections, st.session_state.lang, st.session_state.direction)
    st.divider()
    d1, d2 = st.columns(2)
    docx_bytes = create_docx(st.session_state.cv_sections, st.session_state.lang, st.session_state.direction)
    d1.download_button("📥 Download DOCX", docx_bytes, "Tailored_CV.docx")
    txt_data = "\n".join([f"{SECTION_HEADINGS[st.session_state.lang].get(k, k).upper()}\n{v}\n" for k, v in st.session_state.cv_sections.items() if v.strip() and k != "unclassified"])
    d2.download_button("📥 Download TXT", txt_data, "Tailored_CV.txt")
    st.info("ℹ️ PDF export is temporarily unavailable for better Hebrew support. Please download the Word version and 'Save as PDF'.")
    if st.button("⬅️ Back to Tailoring"): st.session_state.step = 2; st.rerun()
